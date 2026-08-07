"""
CLAUDE-MM8 (Governed Creation, Editing, Review, and Accountable Work
Products): controlled export of a WorkProduct (services/case_workspace.py)
to a real, downloadable file - Section 19's own export contract.

No new dependency: `python-docx` (imported as `docx`, already accepted -
see services/rfi_export.py) covers the narrative-shaped export path;
`openpyxl` (already accepted, MM3) covers the tabular-shaped one.

Two distinct, deliberately un-merged renderers, mirroring rfi_export.py's
own "reuses the same docx-generation library/pattern, not the same
content logic" precedent:
  - build_work_product_docx: narrative artifact_types (report, and any
    other type whose sections read naturally as prose/headed blocks).
  - build_work_product_xlsx: tabular artifact_types (risk_register,
    team_list, and any other type whose sections are naturally rows of a
    consistent shape).

Neither claims perfect round-trip fidelity (Section 19's own explicit
"do not claim perfect round-trip fidelity when it is not guaranteed") -
both are one-way, human-readable renderings of the governed record, not
a serialization format BEEHIVE itself reads back. Re-importing an
exported file is explicitly out of scope this stage (Section 20 concerns
recognizing an ARCHIOSK-created artifact on reopen INSIDE the app, which
`revise_work_product`'s own Supersession-linked draft already provides -
not re-parsing a downloaded .docx/.xlsx back into sections).

Section 27's formula-injection safeguard: any cell value that begins with
a formula-triggering character (=, +, -, @) is prefixed with a leading
apostrophe before being written - the same well-known CSV/XLSX injection
defense every spreadsheet-writing tool needs, applied here so a risk
description a user typed (or an evidence excerpt quoted verbatim) can
never be silently interpreted as a formula by whatever application opens
the exported file.
"""
from __future__ import annotations

import hashlib
import io

import docx
import openpyxl

REMOVED_METADATA_NOTE = (
    "Internal machinery deliberately excluded from this export: raw object "
    "ids, edit history, and governance/audit detail all remain in ARCHIOSK's "
    "own governed record, not in a document meant to be shared."
)

_FORMULA_TRIGGER_CHARS = ("=", "+", "-", "@")


class WorkProductExportError(Exception):
    """Raised when a work product cannot be exported as requested."""


def _sanitize_cell_value(value):
    """Section 27: neutralizes a leading formula-trigger character before
    it ever reaches a spreadsheet cell - a string is returned unchanged
    unless it would otherwise be interpreted as a formula by Excel/Sheets/
    LibreOffice on open."""
    if isinstance(value, str) and value.startswith(_FORMULA_TRIGGER_CHARS):
        return "'" + value
    return value


def _active_sections(work_product: dict) -> list[dict]:
    return sorted(
        (s for s in work_product["sections"] if not s["removed"]),
        key=lambda s: s["order_index"],
    )


def _status_banner(work_product: dict) -> str:
    state = work_product["state"]
    if state == "issued":
        return f"ISSUED — v{work_product['version']} — {work_product.get('issued_at')} by {work_product.get('issued_by')}"
    return f"DRAFT (v{work_product['version']}, state={state}) — not yet issued; for internal review only"


def build_work_product_docx(work_product: dict, sensitivity_note: str | None = None) -> io.BytesIO:
    """Narrative export - one heading per section, content rendered from
    whatever keys the section's own `content` dict carries (a "narrative"
    section uses `text`; any other section_type falls back to a plain
    label: value listing, honest about what it is rather than guessing a
    prose template for structured data it wasn't designed to narrate)."""
    output = docx.Document()
    output.add_heading(work_product["title"], level=1)

    status_p = output.add_paragraph()
    status_p.add_run(_status_banner(work_product)).bold = True

    meta = output.add_paragraph()
    meta.add_run("Project ID: ").bold = True
    meta.add_run(f"{work_product['project_id']}\n")
    meta.add_run("Artifact type: ").bold = True
    meta.add_run(f"{work_product['artifact_type']}\n")
    meta.add_run("Author: ").bold = True
    meta.add_run(f"{work_product['created_by']}\n")
    if sensitivity_note:
        meta.add_run("Sensitivity: ").bold = True
        meta.add_run(f"{sensitivity_note}\n")

    for section in _active_sections(work_product):
        heading_text = section["section_type"].replace("_", " ").title()
        output.add_heading(heading_text, level=2)

        provenance_p = output.add_paragraph()
        provenance_p.add_run("Source: ").italic = True
        provenance_run = provenance_p.add_run(section["content_class"].replace("_", " "))
        provenance_run.italic = True

        if "text" in section["content"]:
            output.add_paragraph(str(section["content"]["text"]))
        else:
            for key, value in section["content"].items():
                p = output.add_paragraph()
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(value))

        if section["evidence_links"]:
            cite_p = output.add_paragraph()
            cite_p.add_run("Cites: ").italic = True
            cite_p.add_run(
                ", ".join(f"{link['object_type']} {link['object_id'][:8]}…" for link in section["evidence_links"])
            )

    footer = output.add_paragraph()
    footer.add_run(REMOVED_METADATA_NOTE).italic = True

    buffer = io.BytesIO()
    output.save(buffer)
    buffer.seek(0)
    return buffer


def build_work_product_xlsx(work_product: dict) -> io.BytesIO:
    """
    Tabular export - one row per active section, columns derived from the
    UNION of every section's own `content` dict keys (in first-seen
    order), so a risk register with a `mitigation` field only on some
    rows still gets one consistent column layout rather than a per-row
    schema. Raises WorkProductExportError if there are no active sections
    to export - an empty spreadsheet would misrepresent an artifact that
    was never actually populated.
    """
    sections = _active_sections(work_product)
    if not sections:
        raise WorkProductExportError("This work product has no active sections to export.")

    columns: list[str] = []
    for section in sections:
        for key in section["content"].keys():
            if key not in columns:
                columns.append(key)

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = work_product["artifact_type"][:31] or "Sheet1"

    sheet.append(["ID", "Type"] + [c.replace("_", " ").title() for c in columns] + ["Source"])
    for cell in sheet[1]:
        cell.font = openpyxl.styles.Font(bold=True)

    for section in sections:
        row = [_sanitize_cell_value(section["id"][:8]), _sanitize_cell_value(section["section_type"])]
        for col in columns:
            row.append(_sanitize_cell_value(section["content"].get(col, "")))
        row.append(_sanitize_cell_value(section["content_class"]))
        sheet.append(row)

    meta_sheet = workbook.create_sheet("Metadata")
    meta_rows = [
        ("Title", work_product["title"]),
        ("Project ID", work_product["project_id"]),
        ("Artifact type", work_product["artifact_type"]),
        ("Status", _status_banner(work_product)),
        ("Author", work_product["created_by"]),
        ("Note", REMOVED_METADATA_NOTE),
    ]
    for label, value in meta_rows:
        meta_sheet.append([_sanitize_cell_value(label), _sanitize_cell_value(value)])

    buffer = io.BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def export_work_product(work_product: dict, export_format: str) -> tuple[io.BytesIO, str]:
    """Dispatches to the correct renderer by format, then computes the
    SHA-256 checksum of the actual exported bytes (Section 19's own
    required export-record field) - the checksum is of what was really
    produced, never a value derived independently that could drift from
    the file a caller actually receives."""
    if export_format == "docx":
        buffer = build_work_product_docx(work_product)
    elif export_format == "xlsx":
        buffer = build_work_product_xlsx(work_product)
    else:
        raise WorkProductExportError(f"Unsupported export format: '{export_format}'. Use 'docx' or 'xlsx'.")

    checksum = hashlib.sha256(buffer.getvalue()).hexdigest()
    buffer.seek(0)
    return buffer, checksum
