"""
Two distinct, deliberately un-merged RFI export mechanisms:

1. build_rfi_docx - the original, project-wide exporter, built from a
   document's flagged cross-requirement contradictions (see
   BHiveParser's consistency-check stage). Unrelated to any specific
   Case or Finding.
2. build_rfi_draft_docx - exports one governed, Case/Finding-scoped
   RFIDraft (services/case_workspace.py) exactly as recorded - never
   fed through (1), which would misrepresent a per-Finding governed
   draft as a project-wide consistency-flag report. Reuses the same
   docx-generation library/pattern, not the same content logic.
"""
from __future__ import annotations

import io

import docx

from services.bhive_parser import ParsedDocument


class RFIExportError(Exception):
    """Raised when there is nothing to export."""


def build_rfi_docx(document: ParsedDocument, operating_environment_label: str | None = None) -> io.BytesIO:
    """
    `operating_environment_label` (CLAUDE-P29), if given, stamps which
    Project Operating Environment this export was produced under --
    optional and additive (existing callers that don't pass it are
    unaffected) since ParsedDocument itself has no operating_environment
    of its own (that lives on ProjectWorkspace, a different store) - the
    caller is expected to look it up and pass the label through.
    """
    if not document.consistency_flags:
        reason = (
            "No flagged contradictions to export."
            if document.consistency_checked
            else "Consistency check hasn't run for this document."
        )
        raise RFIExportError(reason)

    output = docx.Document()

    output.add_heading(f"Request for Information — {document.filename}", level=1)

    meta = output.add_paragraph()
    meta.add_run("Project ID: ").bold = True
    meta.add_run(f"{document.project_id}\n")
    meta.add_run("Ingested: ").bold = True
    meta.add_run(f"{document.ingested_at}\n")
    if operating_environment_label:
        meta.add_run("Project Operating Environment: ").bold = True
        meta.add_run(f"{operating_environment_label}\n")
    meta.add_run("Flagged items: ").bold = True
    meta.add_run(str(len(document.consistency_flags)))

    for i, flag in enumerate(document.consistency_flags, start=1):
        output.add_heading(f"RFI-{i:03d}", level=2)

        p = output.add_paragraph()
        p.add_run("Requirement A: ").bold = True
        p.add_run(flag.requirement_a_text)

        p = output.add_paragraph()
        p.add_run("Requirement B: ").bold = True
        p.add_run(flag.requirement_b_text)

        p = output.add_paragraph()
        p.add_run("Flagged discrepancy: ").bold = True
        p.add_run(flag.explanation)

    buffer = io.BytesIO()
    output.save(buffer)
    buffer.seek(0)
    return buffer


def build_rfi_draft_docx(draft: dict, operating_environment_label: str | None = None) -> io.BytesIO:
    """
    Exports one governed per-Finding RFIDraft (a plain dict, as stored in
    ProjectWorkspace.rfi_drafts) as a professional-facing .docx.

    `operating_environment_label` (CLAUDE-P29): optional, additive, same
    contract as build_rfi_docx's own parameter above.

    Every field below comes directly from the draft's own stored state
    (draft itself, or its own reference_snapshot, captured once at
    creation time - see CaseWorkspaceStore.build_reference_snapshot) -
    nothing is re-derived from live Finding/Case state, and nothing not
    actually present on the record is invented. No formal contractual
    RFI numbering exists in the governed model today - the exported
    identifier is honestly an internal reference (draft.id, truncated),
    not a fabricated sequential RFI number.

    Internal machinery deliberately excluded from the professional
    document: engine_name/engine_version (which analysis engine produced
    the underlying Finding), raw artifact/source/case ids, crop pixel
    coordinates - all of that remains in BEEHIVE's own governed record,
    not in a document meant to be sent to another party.
    """
    snapshot = draft.get("reference_snapshot") or {}
    is_issued = draft.get("status") == "issued"
    short_ref = draft["id"][:8]

    output = docx.Document()
    output.add_heading(f"Request for Information — RFI-{short_ref}", level=1)

    status_p = output.add_paragraph()
    if is_issued:
        status_p.add_run(f"ISSUED — {draft.get('issued_at')} by {draft.get('issued_by')}").bold = True
    else:
        status_p.add_run("DRAFT — not yet issued; for internal review only").bold = True

    meta = output.add_paragraph()
    meta.add_run("Internal reference: ").bold = True
    meta.add_run(f"RFI-{short_ref} (formal RFI numbering not yet implemented)\n")
    meta.add_run("Project ID: ").bold = True
    meta.add_run(f"{draft['project_id']}\n")
    if operating_environment_label:
        meta.add_run("Project Operating Environment: ").bold = True
        meta.add_run(f"{operating_environment_label}\n")
    if snapshot.get("case_title"):
        meta.add_run("Case: ").bold = True
        meta.add_run(f"{snapshot['case_title']}\n")
    if snapshot.get("source_name"):
        meta.add_run("Drawing/Document: ").bold = True
        location = snapshot["source_name"]
        if snapshot.get("page"):
            location += f", page {snapshot['page']}"
        meta.add_run(f"{location}\n")
    if snapshot.get("finding_statement"):
        meta.add_run("Finding: ").bold = True
        meta.add_run(f"{snapshot['finding_statement']}\n")
    if snapshot.get("reviewer_validation"):
        meta.add_run("Reviewer Validation: ").bold = True
        reviewer = f" (reviewed by {snapshot['reviewer']})" if snapshot.get("reviewer") else ""
        meta.add_run(f"{snapshot['reviewer_validation']}{reviewer}")

    output.add_heading("Question", level=2)
    output.add_paragraph(draft.get("question_text") or "(No question text recorded yet.)")

    footer = output.add_paragraph()
    footer.add_run(f"Drafted by {draft.get('created_by')} on {draft.get('created_at')}.")
    if is_issued:
        footer.add_run(f" Issued by {draft.get('issued_by')} on {draft.get('issued_at')}.")

    buffer = io.BytesIO()
    output.save(buffer)
    buffer.seek(0)
    return buffer
