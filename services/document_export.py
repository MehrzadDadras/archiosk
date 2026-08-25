"""
CLAUDE-GO-DOCUMENT-EXPORT-01 - real Word, Excel and PDF files to download.

Product Owner: "My Copilot 365 can create PDF and Excel and Word to download.
Make our app to have equal capabilities."

Two of the three needed nothing new. `python-docx` and `openpyxl` have been
declared dependencies since long before this stage (the parser reads .docx and
.xlsx), and `send_file(as_attachment=True)` is already how RFI drafts leave the
application - see `services/rfi_export.py`'s own `build_rfi_draft_docx`, which
this module deliberately does not duplicate or replace. Only PDF was genuinely
absent: `pypdf` reads and manipulates existing PDFs, it does not compose new
ones. `reportlab` was checked against this project's own constraints with
`tools/dependency_fit.py` before being proposed (PASS on all six: no client
build, flat-JSON storage, no async runtime, no new cloud dependency, no
background worker, Python-native).

WHAT THIS EXPORTS, AND WHAT IT DOES NOT

Only content the project already holds, rendered in a different container. An
export is a VIEW of governed state, never a new assertion about it:

  - nothing here writes to the workspace, and no builder receives the store;
  - nothing is summarised, inferred, ranked or reworded by a model - a
    document that quietly editorialised on the way out would be evidence
    laundering, and a reader has no way to tell the difference;
  - a field's own recorded status travels with it. If a Finding is provisional
    it says so in the file, because a document outlives the screen that
    explained it and will be forwarded to people who never saw that screen.

The callers gate on ACTION_EXPORT (`routes/workspace.py`'s
`_require_export_allowed`) exactly as the RFI export already does. This module
enforces nothing itself and must not be reached around it.
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Optional

# Word and Excel: already declared dependencies, already used elsewhere.
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

XLSX_MIMETYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIMETYPE = "application/pdf"

FORMAT_DOCX = "docx"
FORMAT_XLSX = "xlsx"
FORMAT_PDF = "pdf"
SUPPORTED_FORMATS = (FORMAT_DOCX, FORMAT_XLSX, FORMAT_PDF)

MIMETYPES = {
    FORMAT_DOCX: DOCX_MIMETYPE,
    FORMAT_XLSX: XLSX_MIMETYPE,
    FORMAT_PDF: PDF_MIMETYPE,
}


@dataclass
class ExportTable:
    """One tabular block. `headers` and each row are plain strings already -
    formatting decisions belong to whoever assembled this, not to the writers
    below, so a value cannot be silently reinterpreted on the way into a file."""

    title: str
    headers: list[str]
    rows: list[list[str]] = field(default_factory=list)
    note: Optional[str] = None


@dataclass
class ExportDocument:
    """A whole export, in a container-neutral shape.

    Assembled once and handed to all three writers, so the Word, Excel and PDF
    versions of the same request cannot drift apart in content - only in
    presentation. That is the property worth having: a reviewer who exports the
    same thing twice in two formats must not get two different answers.
    """

    title: str
    subtitle: Optional[str] = None
    # Rendered verbatim, in order, before the tables. Provenance and status
    # lines live here.
    preamble: list[str] = field(default_factory=list)
    tables: list[ExportTable] = field(default_factory=list)


def _safe(value) -> str:
    """Everything reaching a writer is a string. `None` becomes empty rather
    than the word "None", which would otherwise be indistinguishable from a
    field whose real recorded value was the text "None"."""
    if value is None:
        return ""
    return str(value)


# ---------------------------------------------------------------- Excel
def build_xlsx(document: ExportDocument) -> io.BytesIO:
    """One worksheet per table, because a spreadsheet's whole value is that its
    rows can be sorted and filtered - stacking unrelated tables on one sheet
    would take exactly that away."""
    workbook = Workbook()
    workbook.remove(workbook.active)

    if document.preamble:
        sheet = workbook.create_sheet("About")
        sheet["A1"] = document.title
        sheet["A1"].font = Font(bold=True, size=14)
        row_index = 3
        if document.subtitle:
            sheet[f"A{row_index}"] = document.subtitle
            row_index += 1
        for line in document.preamble:
            sheet[f"A{row_index}"] = _safe(line)
            sheet[f"A{row_index}"].alignment = Alignment(wrap_text=True, vertical="top")
            row_index += 1
        sheet.column_dimensions["A"].width = 110

    for table in document.tables:
        # Excel sheet titles are capped at 31 characters and may not contain
        # []:*?/\ - a real Investigation title routinely breaks both rules, and
        # openpyxl raises rather than truncating.
        name = "".join(ch for ch in table.title if ch not in "[]:*?/\\")[:31] or "Sheet"
        sheet = workbook.create_sheet(name)
        sheet.append([_safe(h) for h in table.headers])
        for cell in sheet[1]:
            cell.font = Font(bold=True)
        for row in table.rows:
            sheet.append([_safe(value) for value in row])
        if table.note:
            sheet.append([])
            sheet.append([_safe(table.note)])
        for column in range(1, len(table.headers) + 1):
            widths = [len(_safe(row[column - 1])) for row in table.rows if len(row) >= column]
            longest = max([len(_safe(table.headers[column - 1]))] + widths) if widths else len(
                _safe(table.headers[column - 1])
            )
            sheet.column_dimensions[get_column_letter(column)].width = min(60, max(12, longest + 2))
        sheet.freeze_panes = "A2"

    if not workbook.sheetnames:
        workbook.create_sheet("Empty")["A1"] = "Nothing to export."

    buffer = io.BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------- Word
def build_docx(document: ExportDocument) -> io.BytesIO:
    doc = Document()
    doc.add_heading(document.title, level=1)
    if document.subtitle:
        subtitle = doc.add_paragraph(document.subtitle)
        subtitle.runs[0].italic = True

    for line in document.preamble:
        doc.add_paragraph(_safe(line))

    for table in document.tables:
        doc.add_heading(table.title, level=2)
        if not table.rows:
            doc.add_paragraph("Nothing recorded.")
        else:
            grid = doc.add_table(rows=1, cols=len(table.headers))
            grid.style = "Table Grid"
            for index, header in enumerate(table.headers):
                cell = grid.rows[0].cells[index]
                cell.text = _safe(header)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for row in table.rows:
                cells = grid.add_row().cells
                for index, value in enumerate(row[: len(table.headers)]):
                    cells[index].text = _safe(value)
        if table.note:
            note = doc.add_paragraph(_safe(table.note))
            note.runs[0].font.size = Pt(9)
            note.runs[0].italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------- PDF
def build_pdf(document: ExportDocument) -> io.BytesIO:
    """The one format needing a new dependency. `reportlab`'s platypus layer is
    used rather than the canvas API specifically because it paginates tables
    itself - a hand-positioned canvas would silently drop rows past the first
    page, and a report that quietly loses evidence is worse than no report."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        title=document.title,
    )

    flow = [Paragraph(document.title, styles["Title"])]
    if document.subtitle:
        flow.append(Paragraph(document.subtitle, styles["Italic"]))
    flow.append(Spacer(1, 10))
    for line in document.preamble:
        flow.append(Paragraph(_safe(line), styles["BodyText"]))
        flow.append(Spacer(1, 4))

    for index, table in enumerate(document.tables):
        if index:
            flow.append(PageBreak())
        flow.append(Paragraph(_safe(table.title), styles["Heading2"]))
        flow.append(Spacer(1, 6))
        if not table.rows:
            flow.append(Paragraph("Nothing recorded.", styles["BodyText"]))
            continue
        # Every cell is a Paragraph so long values WRAP. Raw strings in a
        # reportlab Table do not wrap - they overflow the column and print off
        # the page edge, which loses content without any error at all.
        data = [[Paragraph(f"<b>{_safe(h)}</b>", styles["BodyText"]) for h in table.headers]]
        for row in table.rows:
            data.append([
                Paragraph(_safe(value), styles["BodyText"])
                for value in row[: len(table.headers)]
            ])
        grid = Table(data, repeatRows=1, hAlign="LEFT")
        grid.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ]))
        flow.append(grid)
        if table.note:
            flow.append(Spacer(1, 6))
            flow.append(Paragraph(_safe(table.note), styles["Italic"]))

    if len(flow) <= 2:
        flow.append(Paragraph("Nothing to export.", styles["BodyText"]))

    doc.build(flow)
    buffer.seek(0)
    return buffer


BUILDERS = {
    FORMAT_DOCX: build_docx,
    FORMAT_XLSX: build_xlsx,
    FORMAT_PDF: build_pdf,
}


def build(document: ExportDocument, export_format: str) -> io.BytesIO:
    if export_format not in BUILDERS:
        raise ValueError(f"Unsupported export format: {export_format}")
    return BUILDERS[export_format](document)
