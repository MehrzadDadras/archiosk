"""
B-Hive Core Chassis — modular ingestion pipeline for RFP/RFQ documents.

Pipeline stages (each is swappable — see `BHiveParser.STAGES`):
    1. extract    — pull raw text out of the uploaded file (pdf/docx/txt/csv)
    2. segment    — split raw text into candidate requirement chunks
    3. classify   — categorize each chunk against the requirement schema
                    (uses the Anthropic API when ANTHROPIC_API_KEY is set;
                    falls back to a rule-based classifier otherwise so the
                    pipeline still runs in dev/test without a key)
    4. consistency — a single holistic pass over the classified requirements
                    looking for cross-requirement contradictions (e.g. a
                    technical spec that can't physically be satisfied by a
                    scheduled milestone). Requires ANTHROPIC_API_KEY — there
                    is no rule-based fallback, since this needs actual
                    semantic reasoning across lines, not per-line keyword
                    matching. Best-effort: never blocks ingestion, and is
                    honest in its output about whether it actually ran.
    5. assemble   — build the final ParsedDocument record

Each stage is a small, independently testable function/class so new
document types or classification strategies can be added without
touching the others.
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
import time
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# CLAUDE-P19: a real, bump-on-meaningful-change marker for the Golden
# Laboratory regression suite's own "prompt/configuration version"
# tracking - see requirement_investigation.py's INVESTIGATION_PROMPT_
# VERSION for the same convention. Last touched by CLAUDE-P23's fix
# requiring verbatim per-pair evidence and an explicit reconciliation-
# checked field, addressing a controlled-experiment-confirmed order/
# adjacency sensitivity in false-positive rate.
CONSISTENCY_PROMPT_VERSION = "p23"

# How long to wait on a single classification batch before giving up on it.
DEFAULT_CLASSIFY_TIMEOUT_SECONDS = 30.0

# Wall-clock ceiling for the whole classify stage, regardless of how many
# batches a large document needs. Bounds worst-case request time so a big
# upload can't run past the Gunicorn/nginx timeouts one batch at a time —
# see deploy/gunicorn.conf.py and deploy/nginx.conf for the matching values.
DEFAULT_CLASSIFY_TOTAL_BUDGET_SECONDS = 90.0

# Single one-shot call, so its worst case is bounded regardless of document
# size (unlike the classify stage, which scales with chunk count) — but it
# still adds to the same request, so it's budgeted into the same
# Gunicorn/nginx timeouts as the classify stage.
DEFAULT_CONSISTENCY_TIMEOUT_SECONDS = 25.0

# Safety valve so a very large document doesn't blow up the consistency
# prompt's size/cost unboundedly. Not env-configurable — it's an internal
# guard, not something operators need to tune.
DEFAULT_CONSISTENCY_MAX_ITEMS = 150

REQUIREMENT_CATEGORIES = [
    "scope_of_work",
    "technical_specification",
    "compliance_legal",
    "budget_commercial",
    "schedule_milestone",
    "submission_instruction",
    "evaluation_criteria",
    "other",
]


class ParserError(Exception):
    """Raised when a document cannot be parsed into requirement records."""


@dataclass
class RequirementItem:
    id: str
    text: str
    category: str
    confidence: float
    source_line: int


@dataclass
class ConsistencyFlag:
    id: str
    requirement_a_id: str
    requirement_a_text: str
    requirement_b_id: str
    requirement_b_text: str
    explanation: str
    # CLAUDE-P23: forces the model to ground each flag in specific,
    # verbatim text from BOTH sides and to have explicitly checked for a
    # reconciling exception BEFORE the pair can appear in the output at
    # all - see _build_consistency_prompt's own docstring for why. Empty-
    # string/False defaults keep this backward compatible with anything
    # constructing a ConsistencyFlag from an older-shaped response.
    requirement_a_evidence: str = ""
    requirement_b_evidence: str = ""
    reconciliation_checked: bool = False


@dataclass
class ParsedDocument:
    project_id: str
    filename: str
    ingested_at: str
    requirements: list[RequirementItem] = field(default_factory=list)
    milestones: list[dict[str, Any]] = field(default_factory=list)
    # Batch H: the raw structured tables (headers + rows) found in the
    # source document, kept separate from `requirements` - a future
    # reconciliation/arithmetic-check capability (still deferred, not
    # built here) needs the real column values, not the header-labeled
    # text reconstruction that feeds classification below.
    tables: list[dict[str, Any]] = field(default_factory=list)
    consistency_flags: list[ConsistencyFlag] = field(default_factory=list)
    # Distinguishes "checked, found nothing" from "didn't actually check" —
    # e.g. no ANTHROPIC_API_KEY, a timeout, or a malformed model response.
    # consistency_flags being empty on its own can't tell you which happened.
    consistency_checked: bool = False
    consistency_note: str | None = None
    # Where the ORIGINAL uploaded bytes are stored, if at all (services/
    # ingestion.py persists them after a successful parse; a pre-existing
    # ParsedDocument ingested before this field existed simply has None
    # here — an honest gap, not backfilled/fabricated). Lets the
    # professional actually open the source that generated this project's
    # extracted requirements, not just read the extraction's output.
    original_file_path: str | None = None
    original_file_hash: str | None = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "project_id": self.project_id,
            "filename": self.filename,
            "ingested_at": self.ingested_at,
            "requirements": [r.__dict__ for r in self.requirements],
            "milestones": self.milestones,
            "tables": self.tables,
            "consistency_flags": [f.__dict__ for f in self.consistency_flags],
            "consistency_checked": self.consistency_checked,
            "consistency_note": self.consistency_note,
            "original_file_path": self.original_file_path,
            "original_file_hash": self.original_file_hash,
        }


# -- table-aware extraction (Batch H) ---------------------------------------
# The concrete gap the NREOCRC baseline adjudication named: BHiveParser had
# no notion of tables at all, so a GFM pipe table (e.g. the Functional
# Program's per-department area breakdown) was invisible to it beyond
# meaningless per-line fragments - each row segmented alone, with no header
# context, and multi-row numeric content (department subtotals) unreachable
# by anything downstream.

_MD_HEADING_RE = re.compile(r"^#{1,6}\s")
_TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-+:?$")


def _is_table_row_line(line: str) -> bool:
    stripped = line.strip()
    return len(stripped) >= 2 and stripped.startswith("|") and stripped.endswith("|")


def _split_table_row(line: str) -> list[str]:
    """
    Deliberately minimal (Batch H): does not handle an escaped pipe
    ("\\|") within a cell - no real document seen so far needs it, and
    adding it now would be speculative rather than grounded in an actual
    gap.
    """
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(_TABLE_SEPARATOR_CELL_RE.match(cell) for cell in cells)


def extract_markdown_tables(text: str) -> list[dict]:
    """
    A minimal GFM pipe-table parser. Requires each table/data row to both
    START and END with "|" - true of every table in the NREOCRC corpus and
    of standard GFM table style; a table written without the outer pipes
    is a real GFM variant this does not attempt to support.

    Returns one dict per table found, in document order:
    {"start_line", "end_line"} are 1-indexed and inclusive, using the same
    line numbering _segment already produces elsewhere - "start_line" is
    the header row's line, "end_line" is the last data row's line (or the
    separator row's line, for a table with zero data rows). "headers" is
    the header row's cells; "rows" is a list of cell-lists, one per data
    row. A data row with fewer cells than the header is padded with ""
    (never fabricated content) rather than raising - a single malformed
    row should not discard an otherwise-good table's real rows.
    """
    lines = text.splitlines()
    tables: list[dict] = []
    i = 0
    while i < len(lines):
        if (
            _is_table_row_line(lines[i])
            and i + 1 < len(lines)
            and _is_table_row_line(lines[i + 1])
            and _is_separator_row(_split_table_row(lines[i + 1]))
        ):
            start_line = i + 1  # 1-indexed header line
            headers = _split_table_row(lines[i])
            j = i + 2
            rows: list[list[str]] = []
            while j < len(lines) and _is_table_row_line(lines[j]):
                cells = _split_table_row(lines[j])
                if len(cells) < len(headers):
                    cells = cells + [""] * (len(headers) - len(cells))
                elif len(cells) > len(headers):
                    cells = cells[: len(headers)]
                rows.append(cells)
                j += 1
            tables.append({"start_line": start_line, "end_line": j, "headers": headers, "rows": rows})
            i = j
        else:
            i += 1
    return tables


class BHiveParser:
    """Coordinates the extract -> segment -> classify -> assemble pipeline."""

    def __init__(
        self,
        anthropic_api_key: str | None = None,
        model: str | None = None,
        timeout: float | None = None,
        total_budget: float | None = None,
        consistency_timeout: float | None = None,
    ):
        # Falls back to the environment so every module gets the key the
        # same way — never hardcode it, never pass it in from a route directly.
        self.api_key = anthropic_api_key or os.getenv("ANTHROPIC_API_KEY", "")
        self.model = model or os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6")
        self.timeout = timeout or float(
            os.getenv("ANTHROPIC_TIMEOUT_SECONDS", DEFAULT_CLASSIFY_TIMEOUT_SECONDS)
        )
        self.total_budget = total_budget or float(
            os.getenv("ANTHROPIC_CLASSIFY_BUDGET_SECONDS", DEFAULT_CLASSIFY_TOTAL_BUDGET_SECONDS)
        )
        self.consistency_timeout = consistency_timeout or float(
            os.getenv("ANTHROPIC_CONSISTENCY_TIMEOUT_SECONDS", DEFAULT_CONSISTENCY_TIMEOUT_SECONDS)
        )

    # -- public entrypoint -------------------------------------------------
    def parse(self, raw_bytes: bytes, filename: str) -> ParsedDocument:
        text = self._extract(raw_bytes, filename)
        if not text.strip():
            raise ParserError(f"No extractable text found in '{filename}'.")

        chunks, tables = self._segment(text)
        requirements = self._classify(chunks)
        milestones = self._derive_milestones(requirements)
        consistency_flags, consistency_checked, consistency_note = (
            self._check_consistency(requirements)
        )

        return ParsedDocument(
            project_id=str(uuid.uuid4()),
            filename=filename,
            ingested_at=datetime.now(timezone.utc).isoformat(),
            requirements=requirements,
            milestones=milestones,
            tables=tables,
            consistency_flags=consistency_flags,
            consistency_checked=consistency_checked,
            consistency_note=consistency_note,
        )

    # -- stage 1: extract ---------------------------------------------------
    def _extract(self, raw_bytes: bytes, filename: str) -> str:
        ext = Path(filename).suffix.lower()

        # .md is already plain text - the only thing that makes it
        # different from .txt is that _segment (below) now knows how to
        # read its structure (tables, ATX headings) instead of treating
        # every non-trivial line the same way.
        if ext in (".txt", ".csv", ".md"):
            return raw_bytes.decode("utf-8", errors="ignore")

        if ext == ".docx":
            return self._extract_docx(raw_bytes)

        if ext == ".pdf":
            return self._extract_pdf(raw_bytes)

        raise ParserError(f"Unsupported extension for extraction: {ext}")

    @staticmethod
    def _extract_docx(raw_bytes: bytes) -> str:
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ParserError("python-docx is required to parse .docx files.") from exc

        document = docx.Document(io.BytesIO(raw_bytes))
        # Section headings ("1. Scope of Work") aren't requirements, and
        # their short, keyword-heavy text can itself get misclassified
        # (e.g. a "Schedule and Milestones" heading matching the
        # schedule_milestone keyword list and appearing as a fake
        # milestone). docx already tells us which paragraphs are
        # headings via style name — use that instead of guessing.
        body_paragraphs = [
            p.text for p in document.paragraphs
            if not (p.style.name or "").startswith(("Heading", "Title"))
        ]
        return "\n".join(body_paragraphs)

    @staticmethod
    def _extract_pdf(raw_bytes: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ParserError("pypdf is required to parse .pdf files.") from exc

        reader = PdfReader(io.BytesIO(raw_bytes))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    # -- stage 2: segment ----------------------------------------------------
    def _segment(self, text: str) -> tuple[list[tuple[int, str]], list[dict]]:
        """
        Split into non-trivial lines/clauses, keeping 1-indexed line
        numbers, plus (Batch H) any GFM tables found in the text.

        Table-aware: table lines (header/separator/data rows) are
        excluded from the naive per-line pass below - left in, a table
        row would segment as one meaningless raw-pipe fragment with no
        header context. Instead each data row becomes its own
        header-labeled chunk ("Functional Group: ...; Room / Space: ...;
        ..."), giving the existing classify stage legible text, while the
        raw table (headers + rows) is returned separately for a caller
        that wants the real values.

        Also excludes markdown ATX headings ("## 12.1 Standby Power") -
        the same reasoning _extract_docx already applies to Heading/Title
        paragraphs: a heading's short, keyword-heavy text can otherwise
        get misclassified as a real requirement.
        """
        tables = extract_markdown_tables(text)
        table_line_numbers: set[int] = set()
        for table in tables:
            table_line_numbers.update(range(table["start_line"], table["end_line"] + 1))

        lines = text.splitlines()
        chunks: list[tuple[int, str]] = []
        for i, line in enumerate(lines, start=1):
            if i in table_line_numbers:
                continue
            if _MD_HEADING_RE.match(line.strip()):
                continue
            cleaned = line.strip(" \t-*•")
            if len(cleaned) >= 8:
                chunks.append((i, cleaned))

        for table in tables:
            headers = table["headers"]
            data_row_start_line = table["start_line"] + 2  # + header line, + separator line
            for row_index, row in enumerate(table["rows"]):
                labeled = " | ".join(
                    f"{header}: {cell}" for header, cell in zip(headers, row) if header
                )
                if labeled:
                    chunks.append((data_row_start_line + row_index, labeled))

        chunks.sort(key=lambda c: c[0])
        return chunks, tables

    # -- stage 3: classify ----------------------------------------------------
    def _classify(self, chunks: list[tuple[int, str]]) -> list[RequirementItem]:
        if self.api_key:
            try:
                return self._classify_with_model(chunks)
            except Exception:
                # Model classification is best-effort; never let an API hiccup
                # take down ingestion. Fall through to the rule-based path.
                logger.warning(
                    "Model classification failed; falling back to rule-based classification.",
                    exc_info=True,
                )
        return self._classify_with_rules(chunks)

    def _classify_with_model(self, chunks: list[tuple[int, str]]) -> list[RequirementItem]:
        """Batch-classify chunks via the Anthropic API. Requires ANTHROPIC_API_KEY."""
        import anthropic  # imported lazily so the dep is optional in dev

        client = anthropic.Anthropic(api_key=self.api_key, timeout=self.timeout)
        items: list[RequirementItem] = []

        batch_size = 25
        started_at = time.monotonic()
        for start in range(0, len(chunks), batch_size):
            batch = chunks[start:start + batch_size]

            elapsed = time.monotonic() - started_at
            if elapsed > self.total_budget:
                # A large document could otherwise run one 30s-timeout batch
                # after another well past the Gunicorn/nginx request timeout.
                # Once the overall budget is spent, stop calling the model
                # entirely — rule-classify everything that's left in one go.
                remaining = chunks[start:]
                logger.warning(
                    "Classification budget of %.0fs exceeded after %d batch(es); "
                    "falling back to rule-based classification for the "
                    "remaining %d line(s).",
                    self.total_budget, start // batch_size, len(remaining),
                )
                items.extend(self._classify_with_rules(remaining))
                break

            prompt = self._build_classification_prompt(batch)
            try:
                response = client.messages.create(
                    model=self.model,
                    max_tokens=2000,
                    messages=[{"role": "user", "content": prompt}],
                )
            except anthropic.APITimeoutError:
                # Don't let one slow batch discard classification results the
                # model already produced for earlier batches — only this
                # batch's chunks fall back to the rule-based classifier.
                logger.warning(
                    "Anthropic classification request timed out after %.0fs for "
                    "lines %d-%d; falling back to rule-based classification for "
                    "this batch.",
                    self.timeout, batch[0][0], batch[-1][0],
                )
                items.extend(self._classify_with_rules(batch))
                continue

            text_out = "".join(
                block.text for block in response.content if getattr(block, "type", None) == "text"
            )
            items.extend(self._parse_model_output(text_out, batch))

        return items

    @staticmethod
    def _build_classification_prompt(batch: list[tuple[int, str]]) -> str:
        categories = ", ".join(REQUIREMENT_CATEGORIES)
        lines = "\n".join(f"{line_no}: {text}" for line_no, text in batch)
        return (
            "Classify each numbered line from an RFP/RFQ into exactly one of "
            f"these categories: {categories}.\n"
            "Respond ONLY with a JSON array of objects: "
            '[{"line": <int>, "category": "<one of the categories>", '
            '"confidence": <0-1 float>}]. No prose, no markdown fences.\n\n"'
            f"{lines}"
        )

    @staticmethod
    def _parse_model_output(
        text_out: str, batch: list[tuple[int, str]]
    ) -> list[RequirementItem]:
        cleaned = re.sub(r"^```(json)?|```$", "", text_out.strip(), flags=re.MULTILINE).strip()
        try:
            parsed = json.loads(cleaned)
        except json.JSONDecodeError as exc:
            raise ParserError("Model returned non-JSON classification output.") from exc

        text_by_line = dict(batch)
        items = []
        for entry in parsed:
            line_no = entry.get("line")
            if line_no not in text_by_line:
                continue
            items.append(
                RequirementItem(
                    id=str(uuid.uuid4()),
                    text=text_by_line[line_no],
                    category=entry.get("category", "other"),
                    confidence=float(entry.get("confidence", 0.5)),
                    source_line=line_no,
                )
            )
        return items

    @staticmethod
    def _classify_with_rules(chunks: list[tuple[int, str]]) -> list[RequirementItem]:
        """Deterministic fallback classifier — keyword matching, no API needed."""
        # Order matters: more specific phrase-level cues are checked before
        # loose single-word ones so e.g. "evaluated on cost" doesn't get
        # swallowed by the budget_commercial bucket.
        keyword_map = {
            "evaluation_criteria": ("evaluated on", "scoring", "evaluation criteria", "weighted"),
            "schedule_milestone": ("deadline", "milestone", "due by", "completion date"),
            "submission_instruction": ("submit", "proposal must include", "submission"),
            "compliance_legal": ("code", "regulation", "license", "insurance", "liability"),
            "scope_of_work": ("scope", "work shall include", "contractor shall"),
            "technical_specification": ("shall comply with", "specification", "material", "dimension"),
            "budget_commercial": ("budget", "cost", "price", "fee", "$"),
        }

        items = []
        for line_no, text in chunks:
            lowered = text.lower()
            category = "other"
            for cat, keywords in keyword_map.items():
                if any(kw in lowered for kw in keywords):
                    category = cat
                    break
            items.append(
                RequirementItem(
                    id=str(uuid.uuid4()),
                    text=text,
                    category=category,
                    confidence=0.4 if category == "other" else 0.65,
                    source_line=line_no,
                )
            )
        return items

    # -- stage 4: consistency check -------------------------------------------
    def _check_consistency(
        self, requirements: list[RequirementItem], usage_sink: dict | None = None,
    ) -> tuple[list[ConsistencyFlag], bool, str | None]:
        """Best-effort holistic pass for cross-requirement contradictions.

        Returns (flags, checked, note). `checked` is False whenever the
        check didn't actually run (no API key, timeout, bad output) —
        callers must not treat an empty `flags` list alone as "verified
        clean", since that's indistinguishable from "never checked".

        CLAUDE-P23: `usage_sink`, if given a dict, gets populated with
        {"prompt", "raw_response_text", "input_tokens", "output_tokens",
        "latency_seconds", "stop_reason"} - purely additive instrumentation
        for the compounding-suspicion controlled experiment (tools/
        self_test_compounding_suspicion_experiment.py) to preserve exact
        prompts/token usage/latency against the REAL production call path,
        without changing this method's return shape or any existing
        caller's behavior (default None = no change at all).
        """
        if not self.api_key:
            return [], False, "Skipped: no ANTHROPIC_API_KEY configured."

        if len(requirements) < 2:
            return [], False, "Skipped: fewer than two requirements to compare."

        import anthropic  # imported lazily so the dep is optional in dev

        candidates = requirements[:DEFAULT_CONSISTENCY_MAX_ITEMS]
        truncated = len(requirements) > len(candidates)

        client = anthropic.Anthropic(api_key=self.api_key, timeout=self.consistency_timeout)
        prompt = self._build_consistency_prompt(candidates)

        start = time.perf_counter()
        try:
            response = client.messages.create(
                model=self.model,
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt}],
            )
        except anthropic.APITimeoutError:
            logger.warning(
                "Consistency check timed out after %.0fs; skipping.",
                self.consistency_timeout,
            )
            return [], False, "Skipped: request timed out."
        except Exception:
            # Best-effort, like classification — never let this stage take
            # down ingestion.
            logger.warning("Consistency check failed; skipping.", exc_info=True)
            return [], False, "Skipped: an error occurred."
        latency_seconds = time.perf_counter() - start

        text_out = "".join(
            block.text for block in response.content if getattr(block, "type", None) == "text"
        )
        if usage_sink is not None:
            usage_sink["prompt"] = prompt
            usage_sink["raw_response_text"] = text_out
            usage_sink["latency_seconds"] = latency_seconds
            usage_sink["stop_reason"] = response.stop_reason
            usage = getattr(response, "usage", None)
            usage_sink["input_tokens"] = getattr(usage, "input_tokens", None)
            usage_sink["output_tokens"] = getattr(usage, "output_tokens", None)
        cleaned = re.sub(r"^```(json)?|```$", "", text_out.strip(), flags=re.MULTILINE).strip()
        try:
            parsed = json.loads(cleaned)
        except json.JSONDecodeError:
            logger.warning("Consistency check returned non-JSON output; skipping.")
            return [], False, "Skipped: model returned invalid output."

        by_id = {r.id: r for r in candidates}
        flags = []
        for entry in parsed:
            req_a = by_id.get(entry.get("a"))
            req_b = by_id.get(entry.get("b"))
            if not req_a or not req_b:
                continue
            flags.append(
                ConsistencyFlag(
                    id=str(uuid.uuid4()),
                    requirement_a_id=req_a.id,
                    requirement_a_text=req_a.text,
                    requirement_b_id=req_b.id,
                    requirement_b_text=req_b.text,
                    explanation=entry.get("explanation", ""),
                    requirement_a_evidence=entry.get("requirement_a_evidence", ""),
                    requirement_b_evidence=entry.get("requirement_b_evidence", ""),
                    reconciliation_checked=bool(entry.get("reconciliation_checked", False)),
                )
            )

        note = (
            f"Checked the first {len(candidates)} of {len(requirements)} requirements."
            if truncated else None
        )
        return flags, True, note

    @staticmethod
    def _build_consistency_prompt(requirements: list[RequirementItem]) -> str:
        # CLAUDE-P16: extended beyond numeric/schedule contradictions to
        # cover genuinely SEMANTIC/operational conflicts (two individually
        # reasonable obligations that cannot both govern as written, or a
        # constraint that changes what another requirement effectively
        # permits) - a real production gap found while building the
        # semantic-conflict smoke-test tier, fixed here rather than only
        # inside the benchmark, since every real ingestion's consistency
        # check benefits from it. The paraphrase/drift guardrails exist
        # because "different wording" and "different number" are both
        # weaker signals than "different obligation" - this prompt
        # previously had no guardrail against manufacturing a false
        # positive from wording alone, or missing a real one hidden
        # behind matching wording/numbers.
        # CLAUDE-P22: the requirements list used to come AFTER the "Respond
        # ONLY with JSON" instruction - meaning the actual content to reason
        # about was the LAST thing in the prompt, not the formatting
        # instruction. Under a harder, multi-clause operational-contradiction
        # case (found while admission-reviewing a revised candidate
        # specimen), the model reliably started answering in prose before
        # emitting the JSON array, breaking parsing outright. Moving the
        # requirements list before the reasoning guidance and restating the
        # JSON-only instruction as the LAST thing in the prompt is the
        # standard fix for this class of instruction-adherence issue -
        # every real ingestion's consistency check benefits, not just this
        # one specimen.
        # CLAUDE-P23: a controlled experiment (tools/self_test_compounding_
        # suspicion_experiment.py) found that when two requirements about
        # the SAME topic are NOT adjacent in this list (separated by
        # unrelated items), false-positive risk rises sharply - the model's
        # pairwise comparison quality degrades with distance, producing
        # flags grounded in surface features (e.g. a differing party label)
        # rather than the actual obligation. A same-run example also showed
        # the model's OWN prose reasoning conclude "should not be flagged"
        # for a reconciled pair while still including it in the output -
        # reasoning and action diverging within one response. Forcing a
        # verbatim quote from BOTH sides plus an explicit reconciliation-
        # checked flag, as STRUCTURED JSON fields (never free prose, so this
        # doesn't reopen the CLAUDE-P22 parsing issue), makes "did I actually
        # re-anchor on the real text" and "did I check for an exception"
        # explicit per-pair commitments tied directly to inclusion in the
        # output, rather than an implicit judgment call that can silently
        # diverge from what gets reported.
        lines = "\n".join(f"{r.id}: [{r.category}] {r.text}" for r in requirements)
        return (
            "You are reviewing a procurement document's extracted requirements "
            "for internal contradictions.\n\n"
            f"{lines}\n\n"
            "Consider BOTH kinds of contradiction:\n"
            "1. Numeric/schedule/scope contradictions - e.g. a technical "
            "specification that cannot physically be satisfied by a scheduled "
            "milestone deadline, a budget figure that conflicts with the stated "
            "scope, or compliance terms that conflict with the schedule.\n"
            "2. Semantic/operational contradictions - two individually "
            "reasonable requirements that cannot BOTH govern as written (e.g. "
            "one requires a system remain continuously operational while "
            "another requires the same equipment shut down under the same "
            "condition), or a constraint/qualification elsewhere that changes "
            "what a requirement effectively permits in practice (e.g. an "
            "'unrestricted access' requirement narrowed by a separate access-"
            "control requirement) - unless a third requirement explicitly "
            "provides the exception or mechanism that reconciles them, in "
            "which case there is no contradiction to report.\n\n"
            "Three guardrails against false positives:\n"
            "- Different WORDING for the same requirement is not a "
            "contradiction - do not flag two requirements merely because "
            "project-native terminology differs; judge whether the actual "
            "obligation differs, not the vocabulary. This also covers "
            "different PARTY OR ROLE labels (e.g. 'the Contractor' vs 'the "
            "Design-Builder') describing what may be the same real party "
            "under a different name - do not flag a 'duplicative obligation' "
            "or 'which party is responsible' conflict based solely on "
            "differing role labels; that is an assumption you are making, "
            "not evidence given to you. Only flag a role/party mismatch if "
            "the requirements THEMSELVES state or clearly imply these are "
            "genuinely different, simultaneously-obligated parties.\n"
            "- Do not assume two statements are equivalent just because they "
            "share the same number or subject - a requirement can restate a "
            "figure while changing what is actually being measured or "
            "required (e.g. 'maintain operation for 96 hours' is a "
            "performance obligation; 'provide capacity nominally equivalent "
            "to 96 hours' is a design-basis estimate - not the same "
            "obligation despite the shared number). Flag this kind of drift "
            "as a contradiction if it changes what's actually required.\n\n"
            "Before including ANY pair, you MUST be able to quote a SPECIFIC, "
            "VERBATIM phrase from EACH side that actually conflicts, and you "
            "MUST have checked every OTHER requirement given above for a "
            "reconciling exception, qualification, or resolving mechanism. "
            "If you cannot quote specific verbatim text from both sides, or "
            "if a reconciling exception exists anywhere in the given "
            "requirements, DO NOT include the pair - a requirement's own "
            "distance from another in this list is never itself evidence of "
            "a conflict; judge every pair the same way regardless of where "
            "each one appears. Your explanation must describe ONLY what the "
            "two quoted phrases themselves say - if the actual conflict you "
            "want to describe depends on some OTHER fact not present in "
            "either quoted phrase (e.g. which party is named elsewhere, or "
            "context outside those two phrases), that is not a contradiction "
            "between these two phrases and the pair must not be included.\n\n"
            "Respond ONLY with a JSON array of objects, one per contradiction "
            'found: [{"a": "<requirement id>", "b": "<requirement id>", '
            '"requirement_a_evidence": "<verbatim conflicting phrase quoted '
            "from requirement a's own text>\", "
            '"requirement_b_evidence": "<verbatim conflicting phrase quoted '
            "from requirement b's own text>\", "
            '"reconciliation_checked": <true only if you checked every other '
            'given requirement for a reconciling exception and found none>, '
            '"explanation": "<one concrete sentence>"}]. If there are no '
            "contradictions, respond with an empty JSON array: []. No prose, "
            "no reasoning, no markdown fences - the JSON array must be the "
            "entire response."
        )

    # -- stage 5: assemble (milestones) --------------------------------------
    @staticmethod
    def _derive_milestones(requirements: list[RequirementItem]) -> list[dict[str, Any]]:
        milestones = []
        for req in requirements:
            if req.category == "schedule_milestone":
                milestones.append(
                    {
                        "id": str(uuid.uuid4()),
                        "label": req.text[:120],
                        "status": "pending",
                        "source_line": req.source_line,
                    }
                )
        return milestones
