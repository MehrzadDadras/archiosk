"""
CLAUDE-P34/P35 -- Market-Critical Golden-Path MVP Validation, Layer A
(composed automated integration validation).

Proves the P29-P32 mechanisms operate TOGETHER for one project, not
only in isolation -- environment lock, project ownership/access,
capability gating (RFI directionality), Go/No-Go, requirement
promotion/adjudication, a REAL text-document investigation producing a
provisional Finding, ReviewerValidation/Disposition/Apply, per-Finding
RFI export, project-wide RFI export, close-and-reopen persistence, and
cross-user isolation, all in one continuous sequence against one project.

CLAUDE-P35 CORRECTION to this file's own original CLAUDE-P34 docstring:
that earlier version claimed no UI path exists from a text-ingested
RFP/RFQ to a Finding, based on inspecting only the "Analyze this..."
keyword trigger (services/conversation_interpreter.py's _handle_analyze,
which genuinely IS drawing-only). Further investigation found a SEPARATE,
real, working, non-mock path this test now exercises directly:
1. "Discuss this Requirement" (templates/_macros.html's aperture macro,
   rendered next to every governed Requirement) posts to
   routes/workspace.py's discuss_object with anchor_type=requirement -
   at the PROJECT level (no Case open yet), so an investigation-shaped
   question ("Why is this like this?") is honestly declined with a
   "needs_case" offer, never silently run without a Case to hold the
   resulting Finding.
2. Accepting that offer (POST .../apertures/<message_id>/start-
   investigation) creates a real Case and re-runs the SAME original
   text with the SAME anchor now attached to it - THIS is what reaches
   services/conversation_interpreter.py's _handle_investigate_requirement,
   which calls services/requirement_investigation.py's real Anthropic-
   backed investigate_requirement (mocked below, per this repo's
   hermetic-test convention - see CLAUDE.md) and records a genuine
   provisional Finding via the same record_analysis every other Finding
   path uses.
This is a real, discoverable, two-step UI flow (no hidden URL/repository
knowledge required) - not dead code, not a mock, not a gap. The prior
P34 report's "no UI path exists" finding was corrected in the P35 final
report, not silently dropped.

Every ingestion call spies on BHiveParser.parse -- see CLAUDE.md's
hermetic-test convention (the CLAUDE-P31 8.5-hour live-API incident).

Run via:

    python -m unittest discover -s tests -v
"""
from __future__ import annotations

import io
import shutil
import tempfile
import unittest
import uuid
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import patch

import docx
from werkzeug.datastructures import FileStorage
from werkzeug.security import generate_password_hash

from services.bhive_parser import BHiveParser, ParsedDocument
from services.case_workspace import CaseWorkspaceStore
from services.environment_capabilities import DESIGN_BUILDER_PROPONENT
from services.governance import GovernanceLog
from services.ingestion import ingest_upload
from services.requirements_registry import RequirementsRegistry
from services.security_governance import CONTROL_SOURCE_ARCHIOSK_DEFAULT, SecurityGovernanceStore
from services.security_policy import ACTION_EXPORT, DECISION_DENY

# A real-shaped, but here synthetic, construction procurement excerpt --
# labeled synthetic per this stage's own requirement (Part 4). The real-
# document human walkthrough (Layer B) is what exercises an actual
# product-owner-approved specimen; this automated layer must not claim
# to be that.
_SYNTHETIC_RFP_TEXT = (
    b"Section 3 - Scope of Work\n"
    b"The Design-Builder shall provide all labor, materials, and equipment "
    b"required to complete the civic centre renovation.\n\n"
    b"Section 4 - Technical Specification\n"
    b"All structural steel shall conform to CSA G40.21 350W.\n\n"
    b"Section 5 - Schedule\n"
    b"Substantial performance shall be achieved no later than 18 months "
    b"after contract award.\n"
)


def _fake_file(content: bytes, filename: str) -> FileStorage:
    return FileStorage(stream=io.BytesIO(content), filename=filename)


class MarketCriticalGoldenPathTests(unittest.TestCase):
    def setUp(self):
        import app as app_module
        from models import User, db

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_golden_path_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)

        with self.flask_app.app_context():
            db.session.add_all([
                User(username="alice", password_hash=generate_password_hash("x"), role="admin"),
                User(username="carol", password_hash=generate_password_hash("x"), role="read_only"),
            ])
            db.session.commit()

        self.alice = self._client_as("alice", 1, "admin")
        self.carol = self._client_as("carol", 2, "read_only")

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _client_as(self, username, user_id, role):
        client = self.flask_app.test_client()
        with client.session_transaction() as sess:
            sess["user_id"] = user_id
            sess["username"] = username
            sess["role"] = role
        return client

    def _ingest(self, owner: str, project_name: str, content: bytes = _SYNTHETIC_RFP_TEXT):
        """Deterministic, network-independent -- spies on BHiveParser.parse
        directly (never the real classify/consistency-check pipeline),
        per CLAUDE.md's hermetic-test convention."""
        def fake_parse(self_parser, raw_bytes, filename):
            return ParsedDocument(
                project_id=str(uuid.uuid4()), filename=filename,
                ingested_at=datetime.now(timezone.utc).isoformat(), parser_version="test",
                requirements=[],
            )

        with patch.object(BHiveParser, "parse", fake_parse):
            with self.flask_app.app_context():
                return ingest_upload(
                    _fake_file(content, "sample_rfp.txt"), self.flask_app,
                    operating_environment=DESIGN_BUILDER_PROPONENT, owner=owner, project_name=project_name,
                )

    def test_full_composed_golden_path_for_one_project(self):
        # -- 1. Ingest -> project/workspace creation -> environment lock --
        document = self._ingest(owner="alice", project_name="Civic Centre Renovation RFP")
        store = CaseWorkspaceStore(self.tmp_dir)
        workspace = store.get(document.project_id)
        self.assertEqual(workspace.operating_environment, DESIGN_BUILDER_PROPONENT)
        self.assertEqual(workspace.owner, "alice")

        # -- 2. Project access: owner in, unauthorized user out --
        self.assertEqual(self.alice.get(f"/projects/{document.project_id}/workspace").status_code, 200)
        self.assertEqual(self.carol.get(f"/projects/{document.project_id}/workspace").status_code, 404)
        self.assertEqual(self.carol.get(f"/api/v1/documents/{document.project_id}").status_code, 404)

        # -- 3. Allow-list grant -> carol now authorized (composes with
        #    Case-level privacy below: authorized != sees every Case) --
        # Re-fetched immediately before this write: the GET requests
        # above (open_project's own temporal reconciliation, the nav
        # sidebar's backfill check) may have advanced workspace.version
        # on disk since it was first loaded -- every store-layer write
        # in this test re-fetches for the same reason, matching this
        # store's own optimistic-concurrency contract.
        workspace = store.get(document.project_id)
        store.grant_project_access(workspace, username="carol", actor="alice", actor_role="admin")
        self.assertEqual(self.carol.get(f"/projects/{document.project_id}/workspace").status_code, 200)

        # -- 4. Go/No-Go: environment-specific decision vocabulary --
        go_no_go_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "pursue", "decision": "go", "rationale": "Strong fit for our team."},
        )
        self.assertIn(go_no_go_response.status_code, (302, 303))
        workspace = store.get(document.project_id)
        self.assertEqual(len(workspace.go_no_go_assessments), 1)

        # -- 5. Case creation --
        case_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/cases",
            data={"title": "RFP Compliance Review", "objective": "Review scope, spec, and schedule clauses."},
            follow_redirects=True,
        )
        self.assertEqual(case_response.status_code, 200)
        workspace = store.get(document.project_id)
        case = next(c for c in workspace.cases if c["title"] == "RFP Compliance Review")

        # -- 6. Text-document investigative track: promote an extracted
        #    RequirementItem into a governed Requirement, then adjudicate
        #    it. This is the REAL, UI-reachable path for a text RFP --
        #    no drawing, no Finding required. --
        registry = RequirementsRegistry(self.tmp_dir)
        parsed_doc = registry.get(document.project_id)
        from services.bhive_parser import RequirementItem
        parsed_doc.requirements = [
            RequirementItem(id="req-item-1", text="Structural steel shall conform to CSA G40.21 350W.", category="technical_specification", confidence=0.8, source_line=5),
        ]
        registry.save(parsed_doc)

        source_id = next(s["id"] for s in workspace.sources if s["kind"] == "rfq_rfp_document")
        promote_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/cases/{case['id']}/requirement-items/req-item-1/promote",
            data={"source_id": source_id}, follow_redirects=True,
        )
        self.assertEqual(promote_response.status_code, 200)
        workspace = store.get(document.project_id)
        governed_requirement = next(
            r for r in workspace.requirements if r["original_requirement_identifier"] == "req-item-1"
        )

        adjudicate_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/requirements/{governed_requirement['id']}/adjudicate",
            # CLAUDE-POSTCAMEL-COMM-I5A: attribution is now a mandatory,
            # never-defaulted explicit choice at this route.
            data={
                "outcome": "Satisfied", "reasoning": "Structural spec matches CSA G40.21 350W as required.",
                "case_id": case["id"], "attribution": "human_reviewed",
            },
        )
        self.assertIn(adjudicate_response.status_code, (302, 303))
        workspace = store.get(document.project_id)
        self.assertEqual(store.requirement_adjudication_state(workspace, governed_requirement["id"]), "Satisfied")

        # -- 7. Finding/Disposition/Apply/RFI-draft-directionality track --
        # This is the REAL, UI-reachable, two-step text-document
        # investigation path (see module docstring): "Discuss this
        # Requirement" (discuss_object, project-level, no Case open yet)
        # is honestly declined with a needs_case offer, then accepting
        # that offer (start_investigation_from_aperture) opens a real
        # Case and re-runs the same anchored question, this time
        # reaching _handle_investigate_requirement for real.
        # services.requirement_investigation.investigate_requirement is
        # the one Anthropic-API call in this whole path - mocked here
        # per CLAUDE.md's hermetic-test convention, exactly as every
        # other network boundary in this test file already is.
        from services.requirement_investigation import RequirementInvestigationResult

        discuss_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/discuss",
            data={
                "text": "Why is this requirement satisfied - check this against the spec.",
                "anchor_type": "requirement",
                "anchor_id": governed_requirement["id"],
                "anchor_description": governed_requirement["original_requirement_identifier"],
            },
        )
        self.assertIn(discuss_response.status_code, (302, 303))
        workspace = store.get(document.project_id)
        aperture_message = next(
            m for m in workspace.project_conversation
            if m["role"] == "human" and (m.get("anchor") or {}).get("anchor_id") == governed_requirement["id"]
        )

        mocked_result = RequirementInvestigationResult(
            ran=True,
            assessment="The steel grade matches the referenced standard exactly.",
            confidence=0.81,
            supporting_points=["CSA G40.21 350W is explicitly named in Section 4."],
            needs_human_judgment=True,
        )
        with patch("services.conversation_interpreter.investigate_requirement", return_value=mocked_result):
            start_investigation_response = self.alice.post(
                f"/projects/{document.project_id}/workspace/apertures/{aperture_message['id']}/start-investigation",
                follow_redirects=True,
            )
        self.assertEqual(start_investigation_response.status_code, 200)

        workspace = store.get(document.project_id)
        investigation_case = next(
            c for c in workspace.cases
            if c["id"] != case["id"] and c["title"].startswith(governed_requirement["original_requirement_identifier"])
        )
        finding = next(
            f for f in workspace.findings if f["statement"] == mocked_result.assessment
        )
        finding_id = finding["id"]
        self.assertEqual(finding["claim_status"], "provisional")
        self.assertEqual(finding["case_id"], investigation_case["id"])
        case = investigation_case
        store.record_reviewer_validation(workspace, finding_id=finding_id, validation="Correct", reviewer="alice")
        store.record_disposition(workspace, finding_id=finding_id, disposition="Confirmed", reviewer="alice")

        apply_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/cases/{case['id']}/apply", data={"confirm": "once"},
        )
        self.assertIn(apply_response.status_code, (302, 303))
        workspace = store.get(document.project_id)
        applied_finding = next(f for f in workspace.findings if f["id"] == finding_id)
        self.assertEqual(applied_finding["claim_status"], "applied")

        # -- 8. RFI directionality (CLAUDE-P30): Design-Builder/Proponent
        #    may originate; draft + issue through the real routes. --
        draft_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/cases/{case['id']}/rfi-drafts",
            data={"finding_id": finding_id, "question_text": "Please confirm the start trigger for substantial performance."},
        )
        self.assertIn(draft_response.status_code, (302, 303))
        workspace = store.get(document.project_id)
        draft = workspace.rfi_drafts[0]

        issue_response = self.alice.post(
            f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/issue", data={"confirm": "once"},
        )
        self.assertIn(issue_response.status_code, (302, 303))

        # -- 9. Export: per-Finding RFI (direction-stamped) and
        #    project-wide RFI export both reachable and correct. --
        per_finding_export = self.alice.get(f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        self.assertEqual(per_finding_export.status_code, 200)
        exported_text = "\n".join(p.text for p in docx.Document(io.BytesIO(per_finding_export.data)).paragraphs)
        self.assertIn("Design-Builder/Proponent", exported_text)

        # Project-wide export needs a consistency flag on record to have
        # anything to export -- seed one directly (the consistency-check
        # stage itself is out of scope for this composed test; it's
        # exercised elsewhere).
        from services.bhive_parser import ConsistencyFlag
        parsed_doc = registry.get(document.project_id)
        parsed_doc.consistency_checked = True
        parsed_doc.consistency_flags = [
            ConsistencyFlag(id="flag-1", requirement_a_id="a", requirement_a_text="18 months.", requirement_b_id="b", requirement_b_text="12 months.", explanation="Conflicting schedule durations."),
        ]
        registry.save(parsed_doc)
        project_export = self.alice.get(f"/projects/{document.project_id}/workspace/rfi-export")
        self.assertEqual(project_export.status_code, 200)

        # -- 10. Close and reopen: a fresh CaseWorkspaceStore instance
        #    (simulating a new session against the same on-disk state)
        #    must see every governed fact recorded above. --
        fresh_store = CaseWorkspaceStore(self.tmp_dir)
        reloaded = fresh_store.get(document.project_id)
        self.assertEqual(reloaded.operating_environment, DESIGN_BUILDER_PROPONENT)
        self.assertEqual(reloaded.owner, "alice")
        self.assertIn("carol", reloaded.access_allow_list)
        self.assertEqual(len(reloaded.go_no_go_assessments), 1)
        self.assertEqual(
            fresh_store.requirement_adjudication_state(reloaded, governed_requirement["id"]), "Satisfied",
        )
        reloaded_finding = next(f for f in reloaded.findings if f["id"] == finding_id)
        self.assertEqual(reloaded_finding["claim_status"], "applied")
        self.assertEqual(reloaded.rfi_drafts[0]["status"], "issued")

        # -- 11. Unauthorized-user isolation still holds after all of the
        #    above (revoke carol, confirm denial returns, alice unaffected) --
        fresh_store.revoke_project_access(reloaded, username="carol", actor="alice", actor_role="admin")
        self.assertEqual(self.carol.get(f"/projects/{document.project_id}/workspace").status_code, 404)
        self.assertEqual(self.alice.get(f"/projects/{document.project_id}/workspace").status_code, 200)

    def test_security_export_gate_composes_with_project_access(self):
        """CLAUDE-P31's security-policy export gate and CLAUDE-P32's
        project-access gate are independent systems that must both hold
        -- an active DENY baseline blocks even the project's own owner,
        proving one gate cannot silently substitute for the other."""
        document = self._ingest(owner="alice", project_name="Security Gate Composition Check")
        store = CaseWorkspaceStore(self.tmp_dir)
        workspace = store.get(document.project_id)

        security_store = SecurityGovernanceStore(self.tmp_dir)
        record = security_store.get()
        baseline = security_store.create_baseline_draft(record, created_by="alice")
        security_store.add_control_decision(
            record, baseline_id=baseline["id"], action_id=ACTION_EXPORT, decision=DECISION_DENY,
            source_type=CONTROL_SOURCE_ARCHIOSK_DEFAULT, actor="alice",
        )
        security_store.acknowledge_capability_impact(record, baseline["id"], actor="alice")
        security_store.activate_baseline(record, baseline["id"], actor="alice")

        response = self.alice.get(f"/projects/{document.project_id}/workspace/rfi-export", follow_redirects=True)
        self.assertEqual(response.status_code, 200)
        self.assertIn(b"security policy", response.data)

    def test_corrupted_legacy_workspace_does_not_break_listings_or_nav(self):
        """CLAUDE-P34, Part 5: regression coverage for a corrupted
        legacy workspace file (an unrecognized field ProjectWorkspace's
        dataclass shape rejects) -- proves today's code fails closed
        (excludes) rather than crashing every authenticated page when
        one exists alongside real projects. Originally reproduced with
        a real 'reviews' key; CLAUDE-P40-D gave that key a real
        compatibility adapter (CaseWorkspaceStore._hydrate_legacy_
        reviews), so a still-genuinely-unrecognized key is used here
        instead."""
        good_document = self._ingest(owner="alice", project_name="Healthy Project")

        corrupted_project_id = "legacy-corrupted-unrecognized-field"
        registry = RequirementsRegistry(self.tmp_dir)
        registry.save(ParsedDocument(project_id=corrupted_project_id, filename="old.txt", ingested_at="2020-01-01T00:00:00+00:00"))
        corrupted_path = self.tmp_dir / f"{corrupted_project_id}.workspace.json"
        corrupted_path.write_text(
            '{"project_id": "' + corrupted_project_id + '", "totally_unrecognized_field_xyz": []}', encoding="utf-8",
        )

        # The launcher panel and /projects listing must not crash for ANY
        # authenticated user merely because this one corrupted record
        # exists. CLAUDE-P40-E2B1A: Home itself no longer lists Project
        # names (they're the Projects root launcher's own projected
        # children, on /projects only) - checked here is only that Home
        # still renders at all, not that it lists anything.
        home_response = self.alice.get("/")
        self.assertEqual(home_response.status_code, 200)

        projects_response = self.alice.get("/projects")
        self.assertEqual(projects_response.status_code, 200)
        self.assertIn(b"Healthy Project", projects_response.data)


if __name__ == "__main__":
    unittest.main()
