"""
Product Acceleration - Close the RFI / Compliance Loop: end-to-end tests.

Covers the three closely-related product surfaces added this tranche,
all inside the Case Workspace: (1) the authenticated RFI-export route
that reuses the existing services.rfi_export.build_rfi_docx generator
verbatim, (2) the transparent Requirement-compliance rollup derived
from real RequirementAdjudication state (never an invented score), and
(3) the Accepted Knowledge list that closes the Apply-Confirmed-
Findings visual feedback loop. Also re-confirms archive/privacy
protections remain intact after these additions.

Run via:

    python -m unittest discover -s tests -v
"""
from __future__ import annotations

import shutil
import tempfile
import unittest
from pathlib import Path

from services.bhive_parser import ConsistencyFlag, ParsedDocument, RequirementItem
from services.case_workspace import AnalysisTrigger, CaseWorkspaceStore
from services.requirements_registry import RequirementsRegistry


class RFIExportTests(unittest.TestCase):
    def setUp(self):
        import app as app_module

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_rfi_export_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        self.project_id = "test-project-rfi-export"

        self.flag = ConsistencyFlag(
            id="flag-1", requirement_a_id="r-a", requirement_a_text="Contractor shall use ASTM A36 steel.",
            requirement_b_id="r-b", requirement_b_text="Contractor shall use ASTM A992 steel.",
            explanation="Two different steel specifications for the same structural element.",
        )
        self.document = ParsedDocument(
            project_id=self.project_id, filename="rfp.pdf", ingested_at="2026-01-01T00:00:00+00:00",
            consistency_flags=[self.flag], consistency_checked=True,
        )
        RequirementsRegistry(self.tmp_dir).save(self.document)

        self.client = self.flask_app.test_client()
        with self.client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "design-manager"
            sess["role"] = "admin"

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_download_link_reachable_from_case_workspace(self):
        response = self.client.get(f"/projects/{self.project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertIn(f"/projects/{self.project_id}/workspace/rfi-export", body)
        self.assertIn("Download RFI (.docx)", body)

    def test_export_route_returns_a_real_docx(self):
        response = self.client.get(f"/projects/{self.project_id}/workspace/rfi-export")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn(self.project_id, response.headers["Content-Disposition"])
        # A .docx is a real zip archive - the magic bytes prove this is
        # actual generated content, not a stub/placeholder response.
        self.assertEqual(response.data[:2], b"PK")

    def test_exported_content_corresponds_to_the_correct_project(self):
        import io

        import docx

        response = self.client.get(f"/projects/{self.project_id}/workspace/rfi-export")
        document = docx.Document(io.BytesIO(response.data))
        full_text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn(self.project_id, full_text)
        self.assertIn(self.flag.requirement_a_text, full_text)
        self.assertIn(self.flag.requirement_b_text, full_text)

    def test_unauthenticated_request_cannot_download(self):
        anonymous_client = self.flask_app.test_client()
        response = anonymous_client.get(f"/projects/{self.project_id}/workspace/rfi-export", follow_redirects=False)
        self.assertEqual(response.status_code, 302)
        self.assertIn("/login", response.headers["Location"])

    def test_export_of_nonexistent_project_404s(self):
        response = self.client.get("/projects/does-not-exist/workspace/rfi-export")
        self.assertEqual(response.status_code, 404)

    def test_no_export_link_when_nothing_flagged(self):
        clean_project_id = "test-project-clean"
        clean_document = ParsedDocument(
            project_id=clean_project_id, filename="clean.pdf", ingested_at="2026-01-01T00:00:00+00:00",
            consistency_checked=True, consistency_flags=[],
        )
        RequirementsRegistry(self.tmp_dir).save(clean_document)

        response = self.client.get(f"/projects/{clean_project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertNotIn("Download RFI (.docx)", body)
        self.assertIn("no contradictions were flagged", body)


class ComplianceRollupTests(unittest.TestCase):
    def setUp(self):
        import app as app_module

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_compliance_rollup_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        self.project_id = "test-project-compliance"

        self.item_a = RequirementItem(
            id="req-a", text="Contractor shall provide licensed labor.",
            category="compliance_legal", confidence=0.7, source_line=6,
        )
        self.item_b = RequirementItem(
            id="req-b", text="Materials shall comply with ASTM specifications.",
            category="technical_specification", confidence=0.66, source_line=22,
        )
        document = ParsedDocument(
            project_id=self.project_id, filename="rfp.pdf", ingested_at="2026-01-01T00:00:00+00:00",
            requirements=[self.item_a, self.item_b],
        )
        RequirementsRegistry(self.tmp_dir).save(document)

        self.client = self.flask_app.test_client()
        with self.client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "design-manager"
            sess["role"] = "admin"

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _store(self):
        return CaseWorkspaceStore(self.tmp_dir)

    def _create_case(self):
        response = self.client.post(
            f"/projects/{self.project_id}/workspace/cases",
            data={"title": "Investigation", "objective": "x"}, follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        return next(c for c in self._store().get(self.project_id).cases if c["title"] == "Investigation")

    def _rfq_source_id(self):
        return next(s for s in self._store().get(self.project_id).sources if s["kind"] == "rfq_rfp_document")["id"]

    def _promote(self, case_id, item):
        self.client.post(
            f"/projects/{self.project_id}/workspace/cases/{case_id}/requirement-items/{item.id}/promote",
            data={"source_id": self._rfq_source_id()}, follow_redirects=True,
        )
        return next(
            r for r in self._store().get(self.project_id).requirements
            if r["original_requirement_identifier"] == item.id
        )

    def test_rollup_shows_not_yet_assessed_before_any_adjudication(self):
        case = self._create_case()
        self._promote(case["id"], self.item_a)

        response = self.client.get(f"/projects/{self.project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertIn("1 Not Yet Assessed", body)

    def test_rollup_reflects_real_adjudication_states(self):
        case = self._create_case()
        req_a = self._promote(case["id"], self.item_a)
        self._promote(case["id"], self.item_b)

        self.client.post(
            f"/projects/{self.project_id}/workspace/requirements/{req_a['id']}/adjudicate",
            data={"outcome": "Satisfied", "reasoning": "Labor licensure confirmed.", "case_id": case["id"]},
            follow_redirects=True,
        )

        response = self.client.get(f"/projects/{self.project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertIn("1 Satisfied", body)
        self.assertIn("1 Not Yet Assessed", body)

    def test_new_adjudication_updates_the_rollup(self):
        case = self._create_case()
        req_a = self._promote(case["id"], self.item_a)
        req_b = self._promote(case["id"], self.item_b)

        self.client.post(
            f"/projects/{self.project_id}/workspace/requirements/{req_a['id']}/adjudicate",
            data={"outcome": "Satisfied", "reasoning": "Confirmed.", "case_id": case["id"]}, follow_redirects=True,
        )
        self.client.post(
            f"/projects/{self.project_id}/workspace/requirements/{req_b['id']}/adjudicate",
            data={"outcome": "Not Satisfied", "reasoning": "Materials do not comply.", "case_id": case["id"]},
            follow_redirects=True,
        )

        response = self.client.get(f"/projects/{self.project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertIn("1 Satisfied", body)
        self.assertIn("1 Not Satisfied", body)
        self.assertNotIn("Not Yet Assessed", body)

    def test_no_invented_compliance_status_appears(self):
        case = self._create_case()
        req_a = self._promote(case["id"], self.item_a)
        self.client.post(
            f"/projects/{self.project_id}/workspace/requirements/{req_a['id']}/adjudicate",
            data={"outcome": "Satisfied", "reasoning": "Confirmed.", "case_id": case["id"]}, follow_redirects=True,
        )

        response = self.client.get(f"/projects/{self.project_id}/workspace?case={case['id']}")
        body = response.get_data(as_text=True)
        # These would be compliance-shaped values Requirement.status
        # itself explicitly denylists (Prompt 15 #17) - none may ever
        # appear as a rollup label, since the rollup only ever echoes
        # real REQUIREMENT_ADJUDICATION_OUTCOMES / the not-yet-assessed
        # sentinel, never an invented score.
        self.assertNotIn(">Compliant<", body)
        self.assertNotIn(">Non-Compliant<", body)
        self.assertNotIn(">non_compliant<", body)

    def test_rollup_survives_reopen_as_a_fresh_session(self):
        case = self._create_case()
        req_a = self._promote(case["id"], self.item_a)
        self.client.post(
            f"/projects/{self.project_id}/workspace/requirements/{req_a['id']}/adjudicate",
            data={"outcome": "Satisfied", "reasoning": "Confirmed.", "case_id": case["id"]}, follow_redirects=True,
        )

        fresh_client = self.flask_app.test_client()
        with fresh_client.session_transaction() as sess:
            sess["user_id"] = 2
            sess["username"] = "design-manager"
            sess["role"] = "admin"

        response = fresh_client.get(f"/projects/{self.project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertIn("1 Satisfied", body)

    def test_drill_down_link_present(self):
        # CLAUDE-POSTCAMEL-ROOT-I1: Requirements now has its own stable
        # Display surface - Overview's own accordion carries a real
        # cross-page link into it (not a same-page anchor anymore, since
        # the target itself moved off the Overview page), and the
        # governed-requirements anchor lives on that new page.
        case = self._create_case()
        self._promote(case["id"], self.item_a)
        overview = self.client.get(f"/projects/{self.project_id}/workspace?view=overview")
        overview_body = overview.get_data(as_text=True)
        self.assertIn('view=requirements', overview_body)
        requirements_page = self.client.get(f"/projects/{self.project_id}/workspace?view=requirements")
        requirements_body = requirements_page.get_data(as_text=True)
        self.assertIn('id="governed-requirements"', requirements_body)


class ApplyFeedbackTests(unittest.TestCase):
    def setUp(self):
        import app as app_module

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_apply_feedback_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        self.project_id = "test-project-apply-feedback"

        document = ParsedDocument(project_id=self.project_id, filename="rfp.pdf", ingested_at="2026-01-01T00:00:00+00:00")
        RequirementsRegistry(self.tmp_dir).save(document)

        self.client = self.flask_app.test_client()
        with self.client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "design-manager"
            sess["role"] = "admin"

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_accepted_knowledge_visible_after_apply(self):
        store = CaseWorkspaceStore(self.tmp_dir)
        workspace = store.get_or_create(self.project_id)
        source = store.add_source(workspace, name="rfp.pdf", file_path="/tmp/rfp.pdf", kind="owner_project_requirements")
        case = store.create_case(workspace, title="Investigation", objective="x", created_by="design-manager")

        trigger = AnalysisTrigger(trigger_type="user_initiated", triggered_by_actor="design-manager")
        analysis = store.record_analysis(
            workspace, case_id=case["id"], source_ids=[source["id"]], objective="x",
            engine_name="test", engine_version="1.0",
            findings=[{"statement": "Beam undersized per drawing S-101.", "machine_confidence": 0.8, "source_id": source["id"]}],
            trigger=trigger,
        )
        finding_id = analysis["finding_ids"][0]
        store.record_reviewer_validation(workspace, finding_id=finding_id, validation="Correct", reviewer="design-manager")
        store.record_disposition(workspace, finding_id=finding_id, disposition="Confirmed", reviewer="design-manager")

        response = self.client.get(f"/projects/{self.project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertNotIn("View accepted knowledge", body)  # nothing applied yet

        apply_response = self.client.post(
            f"/projects/{self.project_id}/workspace/cases/{case['id']}/apply",
            data={"confirm": "once"}, follow_redirects=True,
        )
        self.assertEqual(apply_response.status_code, 200)

        response = self.client.get(f"/projects/{self.project_id}/workspace?view=overview")
        body = response.get_data(as_text=True)
        self.assertIn("View accepted knowledge", body)
        self.assertIn("Beam undersized per drawing S-101.", body)
        self.assertIn("established by design-manager", body)


class ArchivePrivacyStillIntactTests(unittest.TestCase):
    """Light re-confirmation that this tranche's additions did not
    weaken already-established archive/privacy protections."""

    def setUp(self):
        import app as app_module

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_archive_privacy_recheck_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        self.project_id = "test-project-archive-recheck"

        document = ParsedDocument(project_id=self.project_id, filename="rfp.pdf", ingested_at="2026-01-01T00:00:00+00:00")
        RequirementsRegistry(self.tmp_dir).save(document)

        self.owner_client = self.flask_app.test_client()
        with self.owner_client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "owner1"
            sess["role"] = "read_only"

        self.other_client = self.flask_app.test_client()
        with self.other_client.session_transaction() as sess:
            sess["user_id"] = 2
            sess["username"] = "other-user"
            sess["role"] = "read_only"

        # CLAUDE-P32: see tests/test_case_privacy.py's identical setUp
        # comment.
        store = CaseWorkspaceStore(self.tmp_dir)
        workspace = store.get_or_create(self.project_id)
        store.set_project_owner(workspace, owner="owner1", actor="owner1")
        store.grant_project_access(workspace, username="other-user", actor="owner1", actor_role="read_only")

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_apply_still_rejected_on_archived_case(self):
        store = CaseWorkspaceStore(self.tmp_dir)
        response = self.owner_client.post(
            f"/projects/{self.project_id}/workspace/cases",
            data={"title": "Investigation", "objective": "x"}, follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        workspace = store.get(self.project_id)
        case = next(c for c in workspace.cases if c["title"] == "Investigation")

        self.owner_client.post(f"/projects/{self.project_id}/workspace/cases/{case['id']}/archive", follow_redirects=True)

        apply_response = self.owner_client.post(
            f"/projects/{self.project_id}/workspace/cases/{case['id']}/apply",
            data={"confirm": "once"}, follow_redirects=True,
        )
        self.assertEqual(apply_response.status_code, 200)
        workspace = store.get(self.project_id)
        reloaded = next(c for c in workspace.cases if c["id"] == case["id"])
        self.assertEqual(reloaded["status"], "archived")

    def test_private_case_still_invisible_to_other_user(self):
        store = CaseWorkspaceStore(self.tmp_dir)
        response = self.owner_client.post(
            f"/projects/{self.project_id}/workspace/cases",
            data={"title": "Private one", "objective": "x"}, follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        workspace = store.get(self.project_id)
        case = next(c for c in workspace.cases if c["title"] == "Private one")

        response = self.other_client.get(f"/projects/{self.project_id}/workspace?case={case['id']}")
        body = response.get_data(as_text=True)
        self.assertNotIn("Private one", body)


if __name__ == "__main__":
    unittest.main()
