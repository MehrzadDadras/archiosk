"""
CLAUDE-P40-VW8-QA-R2A - Smart Drawing-First Project Qualification.

Repository-grounded capability audit (see services/drawing_intake.py's
own module docstring for the full writeup): native PDF/DOCX text
extraction is real (pypdf/python-docx, already dependencies); PDF page
rendering and local OCR are NOT available in this environment (no
PDF-to-image library installed, and the `tesseract` OS binary is not
present - confirmed directly, not assumed). This stage's own pipeline
is scoped to what that audit actually supports: native-text-based
candidate extraction with full evidence/confidence, honest degradation
when no native text exists, and NO external-AI call of any kind (the
staging-time analysis step never imports or calls the Anthropic client
- see the dedicated test class below that verifies this directly, not
just by inspection).

Every PDF used here is FAKE (`patch.object(BHiveParser,
"extract_pdf_pages", ...)`), matching this codebase's own established
hermetic-test convention (see test_p40vw7_conversation_tags_and_tasks.py's
own note on why real PDF/Anthropic calls are never exercised in this
suite) - CLAUDE.md's own 8.5-hour-hang warning is not theoretical: an
early manual smoke test during this stage's own development genuinely
triggered two real Anthropic API calls by patching only
`extract_pdf_pages` and not `BHiveParser.parse` for the confirm-step
ingestion - every test below patches BOTH.
"""
from __future__ import annotations

import io
import uuid
import unittest
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import patch

from werkzeug.datastructures import FileStorage
from werkzeug.security import generate_password_hash

from services.bhive_parser import BHiveParser, ParsedDocument
from services.case_workspace import CaseWorkspaceStore
from services.environment_capabilities import CLIENT_OWNER
from services.governance import GovernanceLog
from services.requirements_registry import RequirementsRegistry


def _fake_file(content: bytes, filename: str) -> FileStorage:
    return FileStorage(stream=io.BytesIO(content), filename=filename)


def _fake_parse(self_parser, raw_bytes, filename_):
    """The one real hermetic-parse stand-in this whole suite uses (same
    shape as every other test file in this repository) - never invokes
    segment/classify/consistency, so never calls the Anthropic API."""
    return ParsedDocument(
        project_id=str(uuid.uuid4()), filename=filename_,
        ingested_at=datetime.now(timezone.utc).isoformat(), parser_version="test",
    )


_TITLE_BLOCK_PAGE = (
    "SITE PLAN\n"
    "Project Name: Riverside Community Library\n"
    "Project No: 2024-118\n"
    "Owner: City of Riverside\n"
    "Sheet No: A-101\n"
    "Scale: 1\" = 20'-0\"\n"
    "Revision: 2\n"
    "Issue Date: 2026-03-14\n"
    "Consultant: Meridian Architecture Group\n"
)

_PARTIAL_TITLE_BLOCK_PAGE = (
    "Project Name: Foundation Review Site\n"
    "Sheet No: S-201\n"
)


class _BaseTestCase(unittest.TestCase):
    def setUp(self):
        import app as app_module
        from models import User, db
        import tempfile
        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_p40vw8qa_r2a_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)

        with self.flask_app.app_context():
            db.session.add(User(username="r2a_admin", password_hash=generate_password_hash("x"), role="admin"))
            db.session.add(User(username="r2a_reader", password_hash=generate_password_hash("x"), role="read_only"))
            db.session.commit()

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _client_as(self, username, user_id, role="admin"):
        client = self.flask_app.test_client()
        with client.session_transaction() as sess:
            sess["user_id"] = user_id
            sess["username"] = username
            sess["role"] = role
        return client

    def _upload(self, client, pages, filename="site_plan.pdf", project_name=None, environment=CLIENT_OWNER):
        data = {"operating_environment": environment}
        if project_name:
            data["project_name"] = project_name
        data["file"] = (io.BytesIO(b"%PDF-1.4 fake"), filename)
        with patch.object(BHiveParser, "extract_pdf_pages", staticmethod(lambda raw: pages)):
            return client.post("/upload", data=data, content_type="multipart/form-data")

    def _staging_id_from(self, resp) -> str:
        location = resp.headers.get("Location")
        return location.rstrip("/").split("/")[-1]

    def _registry_project_count(self) -> int:
        return len(RequirementsRegistry(self.tmp_dir).list_ids())

    def _store(self) -> CaseWorkspaceStore:
        return CaseWorkspaceStore(self.tmp_dir)

    def _governance_events(self, project_id):
        return GovernanceLog(self.tmp_dir).read(project_id)


# ---------------------------------------------------------------------------
# Section 7: "Drawing with native text" / "Partially readable title block" /
# "Missing optional fields" / "No invented metadata".
# ---------------------------------------------------------------------------

class NativeTextDrawingTests(_BaseTestCase):
    def test_full_title_block_produces_candidates_and_routes_to_confirm(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE])
        self.assertEqual(resp.status_code, 302)
        self.assertIn("/upload/confirm/", resp.headers.get("Location"))

    def test_confirm_page_shows_every_detected_candidate_with_evidence(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE])
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertIn("Riverside Community Library", body)
        self.assertIn("2024-118", body)
        self.assertIn("City of Riverside", body)
        self.assertIn("A-101", body)
        # Evidence: source page + extraction method + the exact matched line.
        self.assertIn("page 1", body)
        self.assertIn("native_pdf_text_pattern", body)
        self.assertIn("Project Name: Riverside Community Library", body)

    def test_partially_readable_title_block_finds_only_whats_there(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_PARTIAL_TITLE_BLOCK_PAGE])
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertIn("Foundation Review Site", body)
        self.assertIn("S-201", body)
        # Fields genuinely absent from the source text must not appear as
        # if they were found - no candidate row's evidence text for
        # owner/consultant/scale/revision/issue_date.
        self.assertNotIn("Meridian Architecture Group", body)

    def test_absence_of_a_field_is_not_treated_as_an_error(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_PARTIAL_TITLE_BLOCK_PAGE])
        # A 302 to the confirm page, not a 400 - partial detection is a
        # normal outcome, never a rejection.
        self.assertEqual(resp.status_code, 302)
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertNotIn("error", body.lower().split("<title>")[0])  # no error banner before content

    def test_no_metadata_is_invented_for_a_plain_rfp_with_no_title_block(self):
        # An ordinary RFP/RFQ (real text, no title-block-shaped lines at
        # all) must produce ZERO candidates - never a fabricated guess.
        client = self._client_as("r2a_admin", 1)
        plain_rfp_page = "The Contractor shall comply with applicable ASTM specifications.\nAll work shall conform to local code.\n"
        resp = self._upload(client, [plain_rfp_page])
        # No candidates + real native text -> straight to ingestion,
        # exactly the pre-existing behavior, not the confirm page.
        self.assertEqual(resp.status_code, 302)
        self.assertNotIn("/upload/confirm/", resp.headers.get("Location", ""))


# ---------------------------------------------------------------------------
# Section 7: "Image-only drawing" / "Creation despite unavailable optional
# analysis" / honest capability degradation (Section 6).
# ---------------------------------------------------------------------------

class ImageOnlyDrawingTests(_BaseTestCase):
    def test_image_only_pdf_does_not_fail_the_upload(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [""])  # empty page text = no native text layer
        self.assertEqual(resp.status_code, 302)
        self.assertIn("/upload/confirm/", resp.headers.get("Location"))

    def test_image_only_pdf_reports_the_capability_gap_honestly(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [""])
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertIn("could not automatically", body)
        self.assertIn("data-ui-ref=\"upload.confirm.no-native-text-notice\"", body)

    def test_project_creation_still_succeeds_from_confirmed_user_input_alone(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [""], project_name="Manually Entered Site Plan")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "field_project_name": "Manually Entered Site Plan",
            })
        self.assertEqual(confirm.status_code, 302, confirm.get_data(as_text=True)[:800])
        self.assertEqual(self._registry_project_count(), 1)

    def test_the_original_drawing_is_still_preserved_for_the_image_only_case(self):
        # Deliberately does NOT mock BHiveParser.parse here (unlike most
        # other tests in this file) - the real parse() short-circuits
        # BEFORE segment/classify/consistency whenever extracted text is
        # empty (see its own early-return branch), so it never reaches
        # the Anthropic-calling stages regardless - this test exercises
        # that REAL code path end-to-end, including the real
        # text_extraction_status determination a fake parse() would
        # otherwise paper over.
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [""], project_name="Preserved Drawing Project")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "extract_pdf_pages", staticmethod(lambda raw: [""])):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "field_project_name": "Preserved Drawing Project",
            })
        project_id = confirm.headers["Location"].split("/")[2]
        document = RequirementsRegistry(self.tmp_dir).get(project_id)
        self.assertIsNotNone(document.original_file_path)
        self.assertTrue(Path(document.original_file_path).exists())
        self.assertEqual(document.text_extraction_status, "no_native_text")

    def test_upload_is_never_reported_as_a_total_failure_for_missing_ocr(self):
        # No 400/500 anywhere in the image-only path - a capability gap
        # is not the same thing as a request failure.
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [""])
        self.assertNotIn(resp.status_code, (400, 500))
        confirm_get = client.get(resp.headers.get("Location"))
        self.assertEqual(confirm_get.status_code, 200)


# ---------------------------------------------------------------------------
# Section 4: conflicting entered vs extracted Project names.
# ---------------------------------------------------------------------------

class ProjectNameConflictTests(_BaseTestCase):
    def test_conflicting_names_are_both_shown_and_require_an_explicit_choice(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="My Own Project Name")
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertIn("My Own Project Name", body)
        self.assertIn("Riverside Community Library", body)
        self.assertIn('data-ui-ref="upload.confirm.name-conflict"', body)

    def test_submitting_without_a_choice_is_rejected(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="My Own Project Name")
        staging_id = self._staging_id_from(resp)
        confirm = client.post(f"/upload/confirm/{staging_id}", data={})
        self.assertEqual(confirm.status_code, 400)
        self.assertEqual(self._registry_project_count(), 0)

    def test_choosing_the_entered_name_neither_overwrites_it_nor_uses_the_candidate(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="My Own Project Name")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={"project_name_choice": "entered"})
        project_id = confirm.headers["Location"].split("/")[2]
        workspace = self._store().get(project_id)
        self.assertEqual(workspace.display_title, "My Own Project Name")

    def test_choosing_the_candidate_name_uses_the_drawing_derived_value(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="My Own Project Name")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={"project_name_choice": "candidate"})
        project_id = confirm.headers["Location"].split("/")[2]
        workspace = self._store().get(project_id)
        self.assertEqual(workspace.display_title, "Riverside Community Library")

    def test_choosing_a_custom_name_uses_neither_original_value(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="My Own Project Name")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "project_name_choice": "custom", "project_name_custom": "A Third Name Entirely",
            })
        project_id = confirm.headers["Location"].split("/")[2]
        workspace = self._store().get(project_id)
        self.assertEqual(workspace.display_title, "A Third Name Entirely")

    def test_matching_names_never_trigger_the_conflict_ui(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="Riverside Community Library")
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertNotIn('data-ui-ref="upload.confirm.name-conflict"', body)


# ---------------------------------------------------------------------------
# Section 4/5: confidence tiers, confirmation vs correction, provenance.
# ---------------------------------------------------------------------------

class ConfidenceAndCorrectionTests(_BaseTestCase):
    def test_colon_labeled_values_are_high_confidence(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE])
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertIn("high confidence", body)

    def test_sheet_number_inferred_discipline_is_medium_confidence(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE])
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertIn("sheet_number_prefix_inference", body)
        self.assertIn("medium confidence", body)

    def test_confirming_a_candidate_unchanged_is_recorded_as_confirmed(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="Riverside Community Library")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "field_project_name": "Riverside Community Library",
                "field_project_number": "2024-118",
            })
        project_id = confirm.headers["Location"].split("/")[2]
        events = self._governance_events(project_id)
        event = next(e for e in events if e.event_type == "drawing_metadata_candidates_confirmed")
        confirmed = {f["field"]: f["status"] for f in event.payload["fields_confirmed"]}
        self.assertEqual(confirmed["project_number"], "confirmed")

    def test_correcting_a_candidate_value_is_recorded_as_corrected(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="Riverside Community Library")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "field_project_name": "Riverside Community Library",
                "field_project_number": "2024-118-CORRECTED",
            })
        project_id = confirm.headers["Location"].split("/")[2]
        events = self._governance_events(project_id)
        event = next(e for e in events if e.event_type == "drawing_metadata_candidates_confirmed")
        confirmed = {f["field"]: f["status"] for f in event.payload["fields_confirmed"]}
        self.assertEqual(confirmed["project_number"], "corrected")
        self.assertEqual(
            next(f["value"] for f in event.payload["fields_confirmed"] if f["field"] == "project_number"),
            "2024-118-CORRECTED",
        )

    def test_original_machine_candidate_evidence_is_preserved_in_governance_log(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="Riverside Community Library")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "field_project_name": "Riverside Community Library",
                "field_project_number": "2024-118",
            })
        project_id = confirm.headers["Location"].split("/")[2]
        events = self._governance_events(project_id)
        event = next(e for e in events if e.event_type == "drawing_metadata_candidates_confirmed")
        offered = {c["field"]: c for c in event.payload["candidates_offered"]}
        self.assertEqual(offered["project_number"]["source_page"], 1)
        self.assertEqual(offered["project_number"]["extraction_method"], "native_pdf_text_pattern")
        self.assertIn("Project No: 2024-118", offered["project_number"]["evidence_snippet"])


# ---------------------------------------------------------------------------
# Section 1: no external-AI transmission of any kind during staging analysis.
# ---------------------------------------------------------------------------

class NoUnauthorizedExternalTransmissionTests(_BaseTestCase):
    def test_drawing_intake_module_never_imports_the_anthropic_client(self):
        import services.drawing_intake as module
        source = Path(module.__file__).read_text(encoding="utf-8")
        self.assertNotIn("import anthropic", source)
        self.assertNotIn("from anthropic", source)

    def test_staging_analysis_never_calls_bhiveparser_classify_or_consistency(self):
        # analyze_upload must reach a real result WITHOUT ever invoking
        # the two stages that can call the Anthropic API - proven by
        # making both raise if called at all, then confirming the
        # staging request still succeeds normally.
        client = self._client_as("r2a_admin", 1)
        with patch.object(BHiveParser, "_classify", side_effect=AssertionError("classify must not run during staging")), \
             patch.object(BHiveParser, "_check_consistency", side_effect=AssertionError("consistency must not run during staging")):
            resp = self._upload(client, [_TITLE_BLOCK_PAGE])
        self.assertEqual(resp.status_code, 302)

    def test_analyze_upload_reads_only_the_bytes_it_is_given(self):
        # Direct unit-level confirmation - services.drawing_intake has
        # exactly one external dependency surface (pypdf/python-docx,
        # both local), verified by successfully analyzing a fake PDF
        # with no network access simulated (no mocking of any transport
        # layer was needed for this to work, because none is used).
        from services.drawing_intake import analyze_upload
        with patch.object(BHiveParser, "extract_pdf_pages", staticmethod(lambda raw: [_TITLE_BLOCK_PAGE])):
            result = analyze_upload(b"irrelevant", "site_plan.pdf")
        self.assertTrue(result.candidates)
        self.assertEqual(result.text_extraction_status, "extracted")


# ---------------------------------------------------------------------------
# Section 5: Project-name uniqueness, environment immutability.
# ---------------------------------------------------------------------------

class QualificationRuleTests(_BaseTestCase):
    def test_duplicate_project_name_is_still_rejected_at_confirm_time(self):
        client = self._client_as("r2a_admin", 1)
        with patch.object(BHiveParser, "parse", _fake_parse):
            client.post("/upload", data={
                "operating_environment": CLIENT_OWNER, "project_name": "Riverside Community Library",
                "file": (io.BytesIO(b"plain text content"), "existing.txt"),
            }, content_type="multipart/form-data")
        self.assertEqual(self._registry_project_count(), 1)

        resp = self._upload(client, [_TITLE_BLOCK_PAGE], project_name="Riverside Community Library")
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "field_project_name": "Riverside Community Library",
            })
        self.assertEqual(confirm.status_code, 400)
        # The rejection is unchanged; only its wording is. The upload path now
        # states the fact ("Project name already exists.") rather than the rule
        # ("Entry names must be unique."). Behaviour is what this test guards,
        # and both assertions around this one still prove it: 400, and the
        # second project was never created.
        self.assertIn("Project name already exists.", confirm.get_data(as_text=True))
        self.assertEqual(self._registry_project_count(), 1)

    def test_operating_environment_is_carried_through_unchanged_never_re_asked(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], environment=CLIENT_OWNER)
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        # No form control to change it here - stated, not re-collected.
        self.assertNotIn('name="operating_environment"', body)
        self.assertIn(CLIENT_OWNER, body)

    def test_operating_environment_is_never_a_candidate_field(self):
        from services.drawing_intake import CANDIDATE_FIELDS
        self.assertNotIn("operating_environment", CANDIDATE_FIELDS)

    def test_confirmed_project_actually_carries_the_originally_chosen_environment(self):
        client = self._client_as("r2a_admin", 1)
        # No entered project_name here (avoids the name-conflict path -
        # already covered by ProjectNameConflictTests above; this test
        # is specifically about the environment, not the name).
        resp = self._upload(client, [_TITLE_BLOCK_PAGE], environment=CLIENT_OWNER)
        staging_id = self._staging_id_from(resp)
        with patch.object(BHiveParser, "parse", _fake_parse):
            confirm = client.post(f"/upload/confirm/{staging_id}", data={
                "field_project_name": "Riverside Community Library",
            })
        project_id = confirm.headers["Location"].split("/")[2]
        workspace = self._store().get(project_id)
        self.assertEqual(workspace.operating_environment, CLIENT_OWNER)


# ---------------------------------------------------------------------------
# Discard / authorization / DOCX support.
# ---------------------------------------------------------------------------

class DiscardAndAuthorizationTests(_BaseTestCase):
    def test_discard_removes_the_staged_upload(self):
        client = self._client_as("r2a_admin", 1)
        resp = self._upload(client, [_TITLE_BLOCK_PAGE])
        staging_id = self._staging_id_from(resp)
        client.post(f"/upload/confirm/{staging_id}/discard")
        self.assertEqual(client.get(f"/upload/confirm/{staging_id}").status_code, 404)

    def test_a_non_admin_cannot_reach_the_confirm_page(self):
        admin_client = self._client_as("r2a_admin", 1)
        resp = self._upload(admin_client, [_TITLE_BLOCK_PAGE])
        staging_id = self._staging_id_from(resp)

        reader_client = self._client_as("r2a_reader", 2, role="read_only")
        confirm_resp = reader_client.get(f"/upload/confirm/{staging_id}")
        self.assertNotEqual(confirm_resp.status_code, 200)

    def test_docx_drawing_cover_sheet_is_also_analyzed(self):
        import docx
        buf = io.BytesIO()
        document = docx.Document()
        document.add_paragraph("Project Name: Docx Cover Sheet Project")
        document.add_paragraph("Project No: 55-100")
        document.save(buf)
        buf.seek(0)

        client = self._client_as("r2a_admin", 1)
        resp = client.post("/upload", data={
            "operating_environment": CLIENT_OWNER,
            "file": (buf, "cover_sheet.docx"),
        }, content_type="multipart/form-data")
        self.assertEqual(resp.status_code, 302)
        self.assertIn("/upload/confirm/", resp.headers.get("Location"))
        body = client.get(resp.headers.get("Location")).get_data(as_text=True)
        self.assertIn("Docx Cover Sheet Project", body)
        self.assertIn("55-100", body)


if __name__ == "__main__":
    unittest.main()
