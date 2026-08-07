"""
CLAUDE-MM8 (Governed Creation, Editing, Review, and Accountable Work
Products) tests: CaseWorkspaceStore.create_work_product/add_work_product_
section/edit_work_product_section/accept_work_product_section/remove_
work_product_section/reorder_work_product_sections/record_work_product_
review/approve_work_product_for_issue/issue_work_product/revise_work_
product/resolve_work_product_status/stale_evidence_for_work_product/
record_work_product_export, and services/work_product_export.py's real
DOCX/XLSX renderers.

Run via:

    python -m unittest tests.test_mm8_work_products -v
"""
from __future__ import annotations

import shutil
import tempfile
import unittest
from pathlib import Path

import docx
import openpyxl

from services.case_workspace import (
    CaseWorkspaceError,
    CaseWorkspaceStore,
    ConcurrentModificationError,
    WORK_PRODUCT_STATE_APPROVED_FOR_ISSUE,
    WORK_PRODUCT_STATE_DRAFT,
    WORK_PRODUCT_STATE_ISSUED,
    WORK_PRODUCT_STATE_REVIEWED,
    WORK_PRODUCT_STATE_REVISIONS_REQUIRED,
    WORK_PRODUCT_STATE_SUPERSEDED,
)
from services.work_product_export import (
    WorkProductExportError,
    build_work_product_docx,
    build_work_product_xlsx,
    export_work_product,
)


class WorkProductLifecycleTests(unittest.TestCase):
    def setUp(self):
        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_mm8_"))
        self.store = CaseWorkspaceStore(self.tmp_dir)
        self.project_id = "test-project-mm8"
        self.workspace = self.store.get_or_create(self.project_id)

        self.pdf_source = self.store.add_source(
            self.workspace, name="spec.pdf", file_path="unused-pdf", kind="document", actor="tester",
        )
        pdf_reg = self.store.register_pdf_page_structure(
            self.workspace, self.pdf_source["id"], ["Section 4.2: retaining wall FoS 1.5."], actor="tester",
        )
        self.pdf_evidence_id = pdf_reg["evidence_item_ids"][0]

        self.case = self.store.create_case(self.workspace, title="MM8 test case", objective="test", created_by="tester")

        self.other_workspace = self.store.get_or_create("test-project-mm8-other")
        other_source = self.store.add_source(
            self.other_workspace, name="other.pdf", file_path="unused-other", kind="document", actor="tester",
        )
        other_reg = self.store.register_pdf_page_structure(
            self.other_workspace, other_source["id"], ["Unrelated other-project text."], actor="tester",
        )
        self.other_project_evidence_id = other_reg["evidence_item_ids"][0]

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _issued_work_product(self):
        """Helper: a real work product carried all the way to issued."""
        wp = self.store.create_work_product(
            self.workspace, artifact_type="risk_register", title="Risk Register", created_by="tester",
            case_id=self.case["id"],
        )
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="risk",
            content={"description": "Settlement risk", "probability": "medium", "impact": "high"},
            content_class="human_authored", author="tester",
            evidence_links=[{"object_type": "evidence_item", "object_id": self.pdf_evidence_id}],
        )
        self.store.record_work_product_review(self.workspace, wp["id"], reviewer="reviewer1", role="Design Manager", decision="reviewed")
        self.store.approve_work_product_for_issue(self.workspace, wp["id"], actor="reviewer1")
        return self.store.issue_work_product(self.workspace, wp["id"], actor="reviewer1")

    # -- identity / draft-vs-issued state -----------------------------------

    def test_work_product_identity_and_persistence(self):
        wp = self.store.create_work_product(
            self.workspace, artifact_type="report", title="Investigation Summary", created_by="tester",
            case_id=self.case["id"],
        )
        self.assertTrue(wp["id"])
        self.assertEqual(wp["state"], WORK_PRODUCT_STATE_DRAFT)
        self.assertEqual(wp["version"], 1)
        reloaded = self.store.get(self.project_id)
        found = self.store._find(reloaded.work_products, wp["id"])
        self.assertIsNotNone(found)
        self.assertEqual(found["title"], "Investigation Summary")

    def test_falsification_blank_title_rejected(self):
        with self.assertRaises(CaseWorkspaceError):
            self.store.create_work_product(self.workspace, artifact_type="report", title="   ", created_by="tester")

    def test_draft_state_resolution(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        status = self.store.resolve_work_product_status(self.workspace, wp["id"])
        self.assertEqual(status["status"], WORK_PRODUCT_STATE_DRAFT)

    def test_issued_state_resolution_and_checksum(self):
        wp = self._issued_work_product()
        self.assertEqual(wp["state"], WORK_PRODUCT_STATE_ISSUED)
        self.assertIsNotNone(wp["issued_checksum"])
        status = self.store.resolve_work_product_status(self.workspace, wp["id"])
        self.assertEqual(status["status"], WORK_PRODUCT_STATE_ISSUED)
        self.assertEqual(status["issued_checksum"], wp["issued_checksum"])

    def test_status_unresolved_for_unknown_work_product(self):
        status = self.store.resolve_work_product_status(self.workspace, "not-a-real-id")
        self.assertEqual(status["status"], "unresolved")

    # -- evidence insertion / citation preservation / project isolation -----

    def test_evidence_insertion_and_citation_preservation(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="narrative", content={"text": "See spec clause."},
            content_class="direct_evidence_reference", author="tester",
            evidence_links=[{"object_type": "evidence_item", "object_id": self.pdf_evidence_id}],
        )
        self.assertEqual(wp["sections"][0]["evidence_links"][0]["object_id"], self.pdf_evidence_id)

    def test_falsification_unsupported_evidence_link_rejected(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        with self.assertRaises(CaseWorkspaceError):
            self.store.add_work_product_section(
                self.workspace, wp["id"], section_type="narrative", content={"text": "x"},
                content_class="human_authored", author="tester",
                evidence_links=[{"object_type": "evidence_item", "object_id": "ghost-id"}],
            )

    def test_falsification_cross_project_evidence_link_rejected(self):
        """No cross-project artifact linkage (Section 27) - proven the
        same structural way MM6/MM7 already proved it, reusing the exact
        same endpoint validator."""
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        with self.assertRaises(CaseWorkspaceError):
            self.store.add_work_product_section(
                self.workspace, wp["id"], section_type="narrative", content={"text": "x"},
                content_class="human_authored", author="tester",
                evidence_links=[{"object_type": "evidence_item", "object_id": self.other_project_evidence_id}],
            )

    # -- human vs AI provenance / editing -----------------------------------

    def test_ai_proposed_section_transitions_to_edited_on_edit(self):
        """Section 8: 'AI content must not become human-approved simply
        because a user opens the document' - editing an ai_proposed
        section transitions it to edited_ai_proposal, never silently to
        human_authored."""
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="narrative", content={"text": "AI draft text."},
            content_class="ai_proposed", author="ai-engine",
        )
        section_id = wp["sections"][0]["id"]
        wp = self.store.edit_work_product_section(
            self.workspace, wp["id"], section_id, content={"text": "Human-revised text."}, actor="tester",
        )
        self.assertEqual(wp["sections"][0]["content_class"], "edited_ai_proposal")
        self.assertEqual(len(wp["sections"][0]["edit_history"]), 1)
        self.assertEqual(wp["sections"][0]["edit_history"][0]["before"]["text"], "AI draft text.")

    def test_accept_ai_proposal_does_not_rewrite_content_class(self):
        """Acceptance and authorship are deliberately two different
        facts - accepting an AI proposal never changes content_class."""
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="narrative", content={"text": "AI draft."},
            content_class="ai_proposed", author="ai-engine",
        )
        section_id = wp["sections"][0]["id"]
        wp = self.store.accept_work_product_section(self.workspace, wp["id"], section_id, actor="reviewer1")
        self.assertEqual(wp["sections"][0]["content_class"], "ai_proposed")
        self.assertEqual(wp["sections"][0]["accepted_by"], "reviewer1")

    def test_falsification_invalid_content_class_rejected(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        with self.assertRaises(CaseWorkspaceError):
            self.store.add_work_product_section(
                self.workspace, wp["id"], section_type="narrative", content={"text": "x"},
                content_class="not_a_real_class", author="tester",
            )

    # -- remove / reorder -----------------------------------------------------

    def test_remove_section_is_soft_delete(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="narrative", content={"text": "x"},
            content_class="human_authored", author="tester",
        )
        section_id = wp["sections"][0]["id"]
        wp = self.store.remove_work_product_section(self.workspace, wp["id"], section_id, actor="tester", reason="obsolete")
        self.assertTrue(wp["sections"][0]["removed"])
        # Soft-delete only - the record itself is still present, inspectable.
        self.assertEqual(wp["sections"][0]["id"], section_id)

    def test_reorder_sections(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        wp = self.store.add_work_product_section(self.workspace, wp["id"], section_type="narrative", content={"text": "first"}, content_class="human_authored", author="tester")
        wp = self.store.add_work_product_section(self.workspace, wp["id"], section_type="narrative", content={"text": "second"}, content_class="human_authored", author="tester")
        ids = [s["id"] for s in wp["sections"]]
        wp = self.store.reorder_work_product_sections(self.workspace, wp["id"], list(reversed(ids)), actor="tester")
        by_id = {s["id"]: s["order_index"] for s in wp["sections"]}
        self.assertEqual(by_id[ids[0]], 1)
        self.assertEqual(by_id[ids[1]], 0)

    # -- issued-version immutability (falsification) -------------------------

    def test_falsification_cannot_add_section_to_issued_work_product(self):
        wp = self._issued_work_product()
        with self.assertRaises(CaseWorkspaceError):
            self.store.add_work_product_section(
                self.workspace, wp["id"], section_type="narrative", content={"text": "x"},
                content_class="human_authored", author="tester",
            )

    def test_falsification_cannot_edit_section_on_issued_work_product(self):
        wp = self._issued_work_product()
        section_id = wp["sections"][0]["id"]
        with self.assertRaises(CaseWorkspaceError):
            self.store.edit_work_product_section(self.workspace, wp["id"], section_id, content={"description": "tampered"}, actor="tester")

    def test_failed_edit_leaves_issued_content_unchanged(self):
        """The 'failed save' contract: a refused edit never partially
        applies - the issued checksum before and after the refused
        attempt is identical."""
        wp = self._issued_work_product()
        checksum_before = wp["issued_checksum"]
        section_id = wp["sections"][0]["id"]
        try:
            self.store.edit_work_product_section(self.workspace, wp["id"], section_id, content={"description": "tampered"}, actor="tester")
        except CaseWorkspaceError:
            pass
        reloaded = self.store.get_work_product(self.workspace, wp["id"])
        self.assertEqual(reloaded["issued_checksum"], checksum_before)
        self.assertEqual(reloaded["sections"][0]["content"]["description"], "Settlement risk")

    # -- review / approval / issue authority (state machine falsification) --

    def test_falsification_issue_without_approval_rejected(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        with self.assertRaises(CaseWorkspaceError):
            self.store.issue_work_product(self.workspace, wp["id"], actor="tester")

    def test_falsification_approve_without_review_rejected(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        with self.assertRaises(CaseWorkspaceError):
            self.store.approve_work_product_for_issue(self.workspace, wp["id"], actor="tester")

    def test_review_revisions_required_path(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        wp = self.store.record_work_product_review(
            self.workspace, wp["id"], reviewer="reviewer1", role="Design Manager",
            decision="revisions_required", comments="needs more detail",
        )
        self.assertEqual(wp["state"], WORK_PRODUCT_STATE_REVISIONS_REQUIRED)
        self.assertEqual(len(wp["reviews"]), 1)
        self.assertEqual(wp["reviews"][0]["decision"], WORK_PRODUCT_STATE_REVISIONS_REQUIRED)

    def test_falsification_invalid_review_decision_rejected(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        with self.assertRaises(CaseWorkspaceError):
            self.store.record_work_product_review(self.workspace, wp["id"], reviewer="r", role="x", decision="approved")

    # -- revision / correction integrity -------------------------------------

    def test_revise_work_product_preserves_original(self):
        original = self._issued_work_product()
        original_checksum = original["issued_checksum"]

        result = self.store.revise_work_product(self.workspace, original["id"], actor="reviewer1", reason="update mitigation")
        new_wp = result["new_work_product"]

        self.assertNotEqual(new_wp["id"], original["id"])
        self.assertEqual(new_wp["version"], original["version"] + 1)
        self.assertEqual(new_wp["state"], WORK_PRODUCT_STATE_DRAFT)
        # Original untouched, forever recoverable.
        original_reloaded = self.store.get_work_product(self.workspace, original["id"])
        self.assertEqual(original_reloaded["state"], WORK_PRODUCT_STATE_ISSUED)
        self.assertEqual(original_reloaded["issued_checksum"], original_checksum)
        status = self.store.resolve_work_product_status(self.workspace, original["id"])
        self.assertEqual(status["status"], WORK_PRODUCT_STATE_SUPERSEDED)
        self.assertEqual(status["superseded_by_work_product_id"], new_wp["id"])

    def test_revision_sections_are_independent_copies(self):
        """Editing the revision's own copied section must never mutate
        the original issued section by shared reference."""
        original = self._issued_work_product()
        result = self.store.revise_work_product(self.workspace, original["id"], actor="reviewer1")
        new_wp = result["new_work_product"]
        new_section_id = new_wp["sections"][0]["id"]
        self.assertNotEqual(new_section_id, original["sections"][0]["id"])

        self.store.edit_work_product_section(
            self.workspace, new_wp["id"], new_section_id, content={"description": "revised description"}, actor="reviewer1",
        )
        original_reloaded = self.store.get_work_product(self.workspace, original["id"])
        self.assertEqual(original_reloaded["sections"][0]["content"]["description"], "Settlement risk")

    def test_falsification_revise_non_issued_work_product_rejected(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        with self.assertRaises(CaseWorkspaceError):
            self.store.revise_work_product(self.workspace, wp["id"], actor="tester")

    def test_falsification_revise_unknown_work_product_rejected(self):
        with self.assertRaises(CaseWorkspaceError):
            self.store.revise_work_product(self.workspace, "not-a-real-id", actor="tester")

    # -- stale evidence -------------------------------------------------------

    def test_stale_evidence_detected_without_rewriting_issued_content(self):
        wp = self._issued_work_product()
        checksum_before = wp["issued_checksum"]

        raw_source = self.store._find(self.workspace.sources, self.pdf_source["id"])
        raw_source["superseded_by_source_id"] = "irrelevant-successor"
        self.store.save(self.workspace)

        stale = self.store.stale_evidence_for_work_product(self.workspace, wp["id"])
        self.assertTrue(stale["has_stale_or_broken_evidence"])
        reloaded = self.store.get_work_product(self.workspace, wp["id"])
        self.assertEqual(reloaded["issued_checksum"], checksum_before)

    # -- concurrency ------------------------------------------------------------

    def test_concurrent_mutation_protection(self):
        copy_one = self.store.get(self.project_id)
        copy_two = self.store.get(self.project_id)
        self.store.create_work_product(copy_one, artifact_type="report", title="First", created_by="tester")
        with self.assertRaises(ConcurrentModificationError):
            self.store.create_work_product(copy_two, artifact_type="report", title="Second", created_by="tester")

    # -- export / checksum / sensitivity / backward compatibility -----------

    def test_export_checksum_matches_actual_bytes(self):
        import hashlib

        wp = self._issued_work_product()
        buffer, checksum = export_work_product(wp, "xlsx")
        self.assertEqual(hashlib.sha256(buffer.getvalue()).hexdigest(), checksum)

    def test_export_report_docx_content(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="Investigation Report", created_by="tester")
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="narrative", content={"text": "Settlement was observed near the retaining wall."},
            content_class="human_authored", author="tester",
        )
        buffer = build_work_product_docx(wp)
        document = docx.Document(buffer)
        full_text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("Investigation Report", full_text)
        self.assertIn("Settlement was observed", full_text)

    def test_export_risk_register_xlsx_content(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="risk_register", title="Risk Register", created_by="tester")
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="risk",
            content={"description": "Settlement risk", "probability": "medium", "impact": "high"},
            content_class="human_authored", author="tester",
        )
        buffer = build_work_product_xlsx(wp)
        workbook = openpyxl.load_workbook(buffer)
        sheet = workbook.active
        values = [cell.value for row in sheet.iter_rows() for cell in row]
        self.assertIn("Settlement risk", values)

    def test_falsification_export_empty_work_product_rejected(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="risk_register", title="Empty", created_by="tester")
        with self.assertRaises(WorkProductExportError):
            build_work_product_xlsx(wp)

    def test_formula_injection_sanitized_in_xlsx_export(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="risk_register", title="RR", created_by="tester")
        wp = self.store.add_work_product_section(
            self.workspace, wp["id"], section_type="risk",
            content={"description": "=cmd|'/c calc'!A1", "probability": "high", "impact": "high"},
            content_class="human_authored", author="tester",
        )
        buffer = build_work_product_xlsx(wp)
        workbook = openpyxl.load_workbook(buffer)
        sheet = workbook.active
        values = [cell.value for row in sheet.iter_rows(min_row=2) for cell in row if isinstance(cell.value, str)]
        injected = [v for v in values if "cmd" in v]
        self.assertTrue(injected)
        self.assertTrue(all(v.startswith("'") for v in injected))

    def test_record_work_product_export_appends_real_event(self):
        wp = self.store.create_work_product(self.workspace, artifact_type="report", title="R", created_by="tester")
        wp = self.store.record_work_product_export(
            self.workspace, wp["id"], export_format="docx", exported_by="tester", checksum="abc123",
        )
        self.assertEqual(len(wp["exports"]), 1)
        self.assertEqual(wp["exports"][0]["checksum"], "abc123")

    def test_sensitivity_classification_preserved(self):
        wp = self.store.create_work_product(
            self.workspace, artifact_type="report", title="R", created_by="tester",
            sensitivity_classification="confidential",
        )
        reloaded = self.store.get_work_product(self.workspace, wp["id"])
        self.assertEqual(reloaded["sensitivity_classification"], "confidential")

    def test_backward_compatible_with_pre_mm8_rfi_draft(self):
        """The pre-existing RFIDraft mechanism (Section 21's own 'use the
        existing RFI capability as one proof, not the entire product
        direction') remains completely unaffected by anything MM8
        added - the full draft -> issue chain still works exactly as
        before, on the SAME Case a WorkProduct might also be attached
        to."""
        finding_id = self._make_finding()
        self.store.record_reviewer_validation(self.workspace, finding_id=finding_id, validation="Correct", reviewer="tester")
        draft = self.store.create_rfi_draft(self.workspace, finding_id=finding_id, question_text="What is the FoS?", created_by="tester")
        self.assertEqual(draft["status"], "draft")
        issued = self.store.issue_rfi_draft(self.workspace, draft["id"], issued_by="tester")
        self.assertEqual(issued["status"], "issued")

    def _make_finding(self):
        from services.case_workspace import AnalysisTrigger, ANALYSIS_TRIGGER_USER_INITIATED
        trigger = AnalysisTrigger(trigger_type=ANALYSIS_TRIGGER_USER_INITIATED, triggered_by_actor="tester")
        result = self.store.record_analysis(
            self.workspace, source_ids=[self.pdf_source["id"]], objective="test", engine_name="test",
            engine_version="1.0", findings=[{"statement": "A finding", "machine_confidence": 0.8}],
            trigger=trigger, case_id=self.case["id"],
        )
        return result["finding_ids"][0]


class WorkProductRouteTests(unittest.TestCase):
    """Functional tests through the real Flask routes (routes/workspace.py)
    - proves the classic form-POST + redirect + server-rendered-template
    workflow actually renders with real data, not just that the store
    layer is correct in isolation."""

    def setUp(self):
        import app as app_module
        from models import User, db
        from werkzeug.security import generate_password_hash
        from services.bhive_parser import ParsedDocument
        from services.requirements_registry import RequirementsRegistry

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_mm8_routes_"))
        self.project_id = "mm8-route-project"
        document = ParsedDocument(project_id=self.project_id, filename="rfp.md", ingested_at="2026-01-01T00:00:00+00:00")
        RequirementsRegistry(self.tmp_dir).save(document)

        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        with self.flask_app.app_context():
            db.session.add(User(username="mm8admin", password_hash=generate_password_hash("x"), role="admin"))
            db.session.commit()

        self.store = CaseWorkspaceStore(self.tmp_dir)
        self.workspace = self.store.get_or_create(self.project_id)

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _client(self):
        client = self.flask_app.test_client()
        with client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "mm8admin"
            sess["role"] = "admin"
        return client

    def test_create_work_product_route_renders_detail_view(self):
        client = self._client()
        resp = client.post(
            f"/projects/{self.project_id}/workspace/work-products",
            data={"title": "Risk Register", "artifact_type": "risk_register"},
            follow_redirects=True,
        )
        self.assertEqual(resp.status_code, 200)
        body = resp.get_data(as_text=True)
        self.assertIn("Risk Register", body)
        self.assertNotIn("Traceback", body)

    def test_full_lifecycle_through_routes(self):
        client = self._client()
        workspace = self.store.get(self.project_id)
        wp = self.store.create_work_product(workspace, artifact_type="risk_register", title="RR", created_by="mm8admin")
        wp_id = wp["id"]

        resp = client.post(
            f"/projects/{self.project_id}/workspace/work-products/{wp_id}/sections",
            data={"section_type": "risk", "content_class": "human_authored", "description": "Settlement risk",
                  "probability": "medium", "impact": "high"},
            follow_redirects=True,
        )
        self.assertEqual(resp.status_code, 200)
        self.assertIn("Settlement risk", resp.get_data(as_text=True))

        resp = client.post(
            f"/projects/{self.project_id}/workspace/work-products/{wp_id}/review",
            data={"decision": "reviewed", "comments": "looks good"},
            follow_redirects=True,
        )
        self.assertEqual(resp.status_code, 200)

        resp = client.post(
            f"/projects/{self.project_id}/workspace/work-products/{wp_id}/approve-for-issue",
            follow_redirects=True,
        )
        self.assertEqual(resp.status_code, 200)

        # Issue is Approval-Gated - the first POST returns a confirmation
        # page, not a completed issuance (mirrors RFI issue's own
        # precedent exactly).
        resp = client.post(f"/projects/{self.project_id}/workspace/work-products/{wp_id}/issue", follow_redirects=True)
        self.assertEqual(resp.status_code, 200)
        reloaded = self.store.get_work_product(self.store.get(self.project_id), wp_id)
        self.assertNotEqual(reloaded["state"], "issued")

        resp = client.post(
            f"/projects/{self.project_id}/workspace/work-products/{wp_id}/issue",
            data={"confirm": "once"}, follow_redirects=True,
        )
        self.assertEqual(resp.status_code, 200)
        reloaded = self.store.get_work_product(self.store.get(self.project_id), wp_id)
        self.assertEqual(reloaded["state"], "issued")

    def test_export_routes_return_real_files(self):
        client = self._client()
        workspace = self.store.get(self.project_id)
        wp = self.store.create_work_product(workspace, artifact_type="risk_register", title="RR", created_by="mm8admin")
        self.store.add_work_product_section(
            workspace, wp["id"], section_type="risk", content={"description": "x", "probability": "low", "impact": "low"},
            content_class="human_authored", author="mm8admin",
        )

        resp = client.get(f"/projects/{self.project_id}/workspace/work-products/{wp['id']}/export.xlsx")
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(len(resp.data) > 0)

        resp_docx = client.get(f"/projects/{self.project_id}/workspace/work-products/{wp['id']}/export.docx")
        self.assertEqual(resp_docx.status_code, 200)

    def test_unauthenticated_request_rejected(self):
        client = self.flask_app.test_client()
        resp = client.post(
            f"/projects/{self.project_id}/workspace/work-products",
            data={"title": "x", "artifact_type": "report"}, follow_redirects=False,
        )
        self.assertIn(resp.status_code, (302, 401, 403))

    def test_edit_section_form_only_renders_its_own_section_type_fields(self):
        """CLAUDE-POSTCAMEL-P01 regression test: before this stage, the
        "Edit this section" form rendered every field from every section
        type at once (risk description/probability/impact/mitigation/
        owner ALONGSIDE team_member name/role/company/contact ALONGSIDE a
        narrative text box), regardless of the section's own actual
        section_type - a reviewer editing a risk section had no way to
        tell which of the ~10 visible inputs would actually be saved
        (edit_work_product_section only persists the fields
        _WORK_PRODUCT_SECTION_FIELDS maps to that section's own type), so
        filling the wrong group silently discarded that input with zero
        warning. Proves the fix: a risk section's own edit form shows
        risk fields and does NOT show team_member-only fields, and vice
        versa for a team_member section."""
        client = self._client()
        workspace = self.store.get(self.project_id)
        wp = self.store.create_work_product(workspace, artifact_type="report", title="Field Scoping", created_by="mm8admin")
        # add_work_product_section returns the parent WORK PRODUCT dict
        # (not the section) - the section's own id has to be read back
        # from its own sections list, in creation order.
        wp = self.store.add_work_product_section(
            workspace, wp["id"], section_type="risk",
            content={"description": "Settlement risk", "probability": "medium", "impact": "high"},
            content_class="human_authored", author="mm8admin",
        )
        wp = self.store.add_work_product_section(
            self.store.get(self.project_id), wp["id"], section_type="team_member",
            content={"name": "Jane Doe", "role": "Structural Engineer"},
            content_class="human_authored", author="mm8admin",
        )
        risk_section, team_section = wp["sections"][0], wp["sections"][1]
        self.assertEqual(risk_section["section_type"], "risk")
        self.assertEqual(team_section["section_type"], "team_member")

        resp = client.get(f"/projects/{self.project_id}/workspace?work_product={wp['id']}")
        self.assertEqual(resp.status_code, 200)
        body = resp.get_data(as_text=True)

        # The risk section's own edit form carries risk-shaped field
        # names/values...
        self.assertIn('name="probability" placeholder="Probability" value="medium"', body)
        self.assertIn('name="impact" placeholder="Impact" value="high"', body)
        # ...but never a team_member-only field name anywhere near it -
        # a blank `name="role"` input would previously appear directly
        # inside the SAME risk section's edit form.
        risk_form_start = body.index(f'work-products/{wp["id"]}/sections/{risk_section["id"]}"')
        team_form_start = body.index(f'work-products/{wp["id"]}/sections/{team_section["id"]}"')
        risk_form_slice = body[risk_form_start:team_form_start]
        self.assertNotIn('name="role"', risk_form_slice)
        self.assertNotIn('name="company"', risk_form_slice)
        self.assertNotIn('name="contact"', risk_form_slice)
        self.assertNotIn('placeholder="Text"', risk_form_slice)

        # The team_member section's own edit form is the mirror case -
        # role/company/contact with real values, no risk fields. Bounded
        # at "Add a section" - the generic add-new-section form further
        # down the page legitimately contains every field group at once
        # (see that form's own template comment), so an unbounded slice
        # to end-of-body would false-negative against THIS assertion.
        team_form_slice = body[team_form_start:body.index("Add a section")]
        self.assertIn('name="role" placeholder="Role" value="Structural Engineer"', team_form_slice)
        self.assertNotIn('name="probability"', team_form_slice)
        self.assertNotIn('name="mitigation"', team_form_slice)


if __name__ == "__main__":
    unittest.main()
