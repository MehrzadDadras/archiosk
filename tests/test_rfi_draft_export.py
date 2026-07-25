"""
Per-Finding governed RFIDraft export - end-to-end workflow tests.

Exercises the missing output this tranche builds: an individual
governed RFIDraft (Finding -> Validation -> Disposition -> Draft ->
Edit -> Issue) becoming an actual downloadable professional document -
via services.rfi_export.build_rfi_draft_docx and
routes/workspace.py's export_rfi_draft. Deliberately distinct from the
older project-wide consistency-flag exporter (build_rfi_docx /
workspace.export_rfi), covered separately in
tests/test_rfi_compliance_workflow.py.

Run via:

    python -m unittest discover -s tests -v
"""
from __future__ import annotations

import io
import shutil
import tempfile
import unittest
from pathlib import Path

import docx

from services.bhive_parser import ParsedDocument
from services.case_workspace import AnalysisTrigger, CaseWorkspaceStore
from services.requirements_registry import RequirementsRegistry


class RFIDraftExportTests(unittest.TestCase):
    def setUp(self):
        import app as app_module

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_rfi_draft_export_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        self.project_id = "test-project-rfi-draft"

        document = ParsedDocument(project_id=self.project_id, filename="rfp.md", ingested_at="2026-01-01T00:00:00+00:00")
        RequirementsRegistry(self.tmp_dir).save(document)

        self.owner_client = self.flask_app.test_client()
        with self.owner_client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "owner1"
            sess["role"] = "read_only"

        self.other_client = self.flask_app.test_client()
        with self.other_client.session_transaction() as sess:
            sess["user_id"] = 2
            sess["username"] = "other-user"
            sess["role"] = "read_only"

        self.case = self._create_case(self.owner_client)
        self.finding_id = self._create_validated_finding(self.case["id"], "Beam undersized per drawing S-101.")

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _store(self):
        return CaseWorkspaceStore(self.tmp_dir)

    def _create_case(self, client, title="Investigation"):
        response = client.post(
            f"/projects/{self.project_id}/workspace/cases",
            data={"title": title, "objective": "x"}, follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        return next(c for c in self._store().get(self.project_id).cases if c["title"] == title)

    def _rfq_source_id(self, project_id=None):
        return next(
            s for s in self._store().get(project_id or self.project_id).sources if s["kind"] == "rfq_rfp_document"
        )["id"]

    def _create_validated_finding(self, case_id, statement, project_id=None):
        store = self._store()
        pid = project_id or self.project_id
        workspace = store.get(pid)
        trigger = AnalysisTrigger(trigger_type="user_initiated", triggered_by_actor="owner1")
        analysis = store.record_analysis(
            workspace, case_id=case_id, source_ids=[self._rfq_source_id(pid)], objective="x",
            engine_name="test", engine_version="1.0",
            findings=[{"statement": statement, "machine_confidence": 0.7, "source_id": self._rfq_source_id(pid)}],
            trigger=trigger,
        )
        finding_id = analysis["finding_ids"][0]
        workspace = store.get(pid)
        store.record_reviewer_validation(workspace, finding_id=finding_id, validation="Correct", reviewer="owner1")
        return finding_id

    def _create_draft(self, client, case_id, finding_id, question_text="Please clarify beam sizing."):
        response = client.post(
            f"/projects/{self.project_id}/workspace/cases/{case_id}/rfi-drafts",
            data={"finding_id": finding_id, "question_text": question_text}, follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        workspace = self._store().get(self.project_id)
        return next(d for d in workspace.rfi_drafts if d["finding_id"] == finding_id and d["question_text"] == question_text)

    # -- create / edit / export draft -----------------------------------

    def test_create_draft_from_real_finding(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        self.assertEqual(draft["finding_id"], self.finding_id)
        self.assertEqual(draft["status"], "draft")

    def test_edit_question_persists(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        response = self.owner_client.post(
            f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/question",
            data={"question_text": "Updated question text.", "case_id": self.case["id"]}, follow_redirects=True,
        )
        self.assertEqual(response.status_code, 200)
        reloaded = next(d for d in self._store().get(self.project_id).rfi_drafts if d["id"] == draft["id"])
        self.assertEqual(reloaded["question_text"], "Updated question text.")

    def test_download_link_reachable_from_workspace(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        response = self.owner_client.get(f"/projects/{self.project_id}/workspace?case={self.case['id']}")
        body = response.get_data(as_text=True)
        self.assertIn(f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/export", body)
        self.assertIn("Download Draft RFI (.docx)", body)

    def test_export_draft_returns_real_docx_marked_draft(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        response = self.owner_client.get(f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data[:2], b"PK")

        exported = docx.Document(io.BytesIO(response.data))
        full_text = "\n".join(p.text for p in exported.paragraphs)
        self.assertIn("DRAFT", full_text)
        self.assertNotIn("ISSUED", full_text)

    def test_exported_content_corresponds_to_the_exact_draft(self):
        draft = self._create_draft(
            self.owner_client, self.case["id"], self.finding_id, question_text="Please confirm beam size per S-101.",
        )
        response = self.owner_client.get(f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        exported = docx.Document(io.BytesIO(response.data))
        full_text = "\n".join(p.text for p in exported.paragraphs)

        self.assertIn(self.project_id, full_text)
        self.assertIn(self.case["title"], full_text)
        self.assertIn("Beam undersized per drawing S-101.", full_text)
        self.assertIn("Please confirm beam size per S-101.", full_text)
        self.assertIn(draft["id"][:8], full_text)

    def test_another_draft_cannot_be_substituted(self):
        second_finding_id = self._create_validated_finding(self.case["id"], "Column reinforcement unclear on S-102.")
        draft_a = self._create_draft(self.owner_client, self.case["id"], self.finding_id, question_text="Question about beam.")
        draft_b = self._create_draft(self.owner_client, self.case["id"], second_finding_id, question_text="Question about column.")

        response = self.owner_client.get(f"/projects/{self.project_id}/workspace/rfi-drafts/{draft_a['id']}/export")
        full_text = "\n".join(p.text for p in docx.Document(io.BytesIO(response.data)).paragraphs)
        self.assertIn("Question about beam.", full_text)
        self.assertNotIn("Question about column.", full_text)
        self.assertNotIn(draft_b["id"][:8], full_text)

    # -- issue / export issued -----------------------------------------

    def test_issue_and_export_issued_version(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)

        issue_response = self.owner_client.post(
            f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/issue",
            data={"confirm": "once", "case_id": self.case["id"]}, follow_redirects=True,
        )
        self.assertEqual(issue_response.status_code, 200)

        response = self.owner_client.get(f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        self.assertEqual(response.status_code, 200)
        full_text = "\n".join(p.text for p in docx.Document(io.BytesIO(response.data)).paragraphs)
        self.assertIn("ISSUED", full_text)
        self.assertIn("owner1", full_text)

    def test_download_link_says_issued_after_issuance(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        self.owner_client.post(
            f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/issue",
            data={"confirm": "once", "case_id": self.case["id"]}, follow_redirects=True,
        )
        response = self.owner_client.get(f"/projects/{self.project_id}/workspace?case={self.case['id']}")
        body = response.get_data(as_text=True)
        self.assertIn("Download Issued RFI (.docx)", body)
        self.assertNotIn("Download Draft RFI (.docx)", body)

    # -- privacy / authorization -----------------------------------------

    def test_unauthorized_user_cannot_export_private_case_draft(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        response = self.other_client.get(f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        self.assertEqual(response.status_code, 404)

    def test_cross_project_draft_id_reuse_fails(self):
        other_project_id = "test-project-rfi-draft-other"
        RequirementsRegistry(self.tmp_dir).save(
            ParsedDocument(project_id=other_project_id, filename="other.md", ingested_at="2026-01-01T00:00:00+00:00")
        )
        self.owner_client.get(f"/projects/{other_project_id}/workspace")  # seed the workspace file

        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        response = self.owner_client.get(f"/projects/{other_project_id}/workspace/rfi-drafts/{draft['id']}/export")
        self.assertEqual(response.status_code, 404)

    def test_unauthenticated_request_cannot_export(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        anonymous_client = self.flask_app.test_client()
        response = anonymous_client.get(
            f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/export", follow_redirects=False,
        )
        self.assertEqual(response.status_code, 302)
        self.assertIn("/login", response.headers["Location"])

    # -- archive -------------------------------------------------------------

    def test_archived_issued_rfi_remains_downloadable(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        self.owner_client.post(
            f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/issue",
            data={"confirm": "once", "case_id": self.case["id"]}, follow_redirects=True,
        )
        self.owner_client.post(f"/projects/{self.project_id}/workspace/cases/{self.case['id']}/archive", follow_redirects=True)

        response = self.owner_client.get(f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        self.assertEqual(response.status_code, 200)
        full_text = "\n".join(p.text for p in docx.Document(io.BytesIO(response.data)).paragraphs)
        self.assertIn("ISSUED", full_text)

    def test_archived_rfi_cannot_be_mutated_through_edit_or_issue(self):
        draft = self._create_draft(self.owner_client, self.case["id"], self.finding_id)
        self.owner_client.post(f"/projects/{self.project_id}/workspace/cases/{self.case['id']}/archive", follow_redirects=True)

        update_response = self.owner_client.post(
            f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/question",
            data={"question_text": "trying to edit after archive", "case_id": self.case["id"]}, follow_redirects=True,
        )
        self.assertEqual(update_response.status_code, 200)

        issue_response = self.owner_client.post(
            f"/projects/{self.project_id}/workspace/rfi-drafts/{draft['id']}/issue",
            data={"confirm": "once", "case_id": self.case["id"]}, follow_redirects=True,
        )
        self.assertEqual(issue_response.status_code, 200)

        reloaded = next(d for d in self._store().get(self.project_id).rfi_drafts if d["id"] == draft["id"])
        self.assertEqual(reloaded["question_text"], draft["question_text"])
        self.assertEqual(reloaded["status"], "draft")


if __name__ == "__main__":
    unittest.main()
