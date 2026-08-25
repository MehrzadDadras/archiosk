"""CLAUDE-GO-DOCUMENT-EXPORT-01 - real Word, Excel and PDF downloads.

Product Owner: "My Copilot 365 can create PDF and Excel and Word to download.
Make our app to have equal capabilities."

Two of the three needed no new dependency at all - `python-docx` and `openpyxl`
have been declared since long before this stage, and `send_file(as_attachment)`
is already how RFI drafts leave the application. Only PDF was genuinely absent:
`pypdf` reads and manipulates existing PDFs, it cannot compose a laid-out one.

These tests open the produced bytes with real readers rather than checking a
status code. A route that returns 200 and a corrupt file is the failure mode
that matters here, and it is invisible to any assertion about the response
alone.
"""
from __future__ import annotations

import io
import tempfile
import unittest
import zipfile
from pathlib import Path

from werkzeug.security import generate_password_hash

from services.bhive_parser import ParsedDocument, RequirementItem
from services.case_workspace import CaseWorkspaceStore
from services.document_export import (
    SUPPORTED_FORMATS,
    ExportDocument,
    ExportTable,
    build,
)
from services.requirements_registry import RequirementsRegistry

ROOT = Path(__file__).resolve().parents[1]


def _code_of(path: Path) -> str:
    """Strip docstrings and comments before scanning.

    Every guard in this file describes a boundary in prose immediately beside
    the code that honours it, so a bare substring check is satisfied by the
    explanation rather than the behaviour. That has caught this session out
    repeatedly.
    """
    import ast

    source = path.read_text(encoding="utf-8")
    tree = ast.parse(source)
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef, ast.Module)):
            if (node.body and isinstance(node.body[0], ast.Expr)
                    and isinstance(node.body[0].value, ast.Constant)
                    and isinstance(node.body[0].value.value, str)):
                node.body = node.body[1:] or [ast.Pass()]
    return ast.unparse(tree)


def _sample() -> ExportDocument:
    return ExportDocument(
        title="Project Smoke Detector - Findings",
        subtitle="Project psd-001",
        preamble=["Exported from ARCHIOSK.", "Each row carries its own recorded status."],
        tables=[ExportTable(
            "Findings", ["ID", "Statement", "Status"],
            [["F-1", "Damper clearance not dimensioned on A-101", "provisional"],
             ["F-2", "Corridor smoke barrier rating unstated", "provisional"]],
            note="Provisional findings are not conclusions.",
        )],
    )


class TheFilesAreRealTests(unittest.TestCase):
    """Opened with real readers. A route returning 200 and a corrupt file is
    the failure that matters, and no assertion about the response would see
    it."""

    def test_the_word_file_opens_and_carries_the_content(self):
        from docx import Document

        doc = Document(build(_sample(), "docx"))
        text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        self.assertIn("Project Smoke Detector", text + table_text)
        self.assertIn("Damper clearance not dimensioned on A-101", table_text)
        self.assertIn("provisional", table_text)

    def test_the_excel_file_opens_and_carries_the_content(self):
        from openpyxl import load_workbook

        book = load_workbook(build(_sample(), "xlsx"))
        self.assertIn("Findings", book.sheetnames)
        values = [
            str(cell.value)
            for row in book["Findings"].iter_rows()
            for cell in row
            if cell.value is not None
        ]
        self.assertIn("Damper clearance not dimensioned on A-101", values)
        self.assertIn("provisional", values)

    def test_the_pdf_is_a_real_pdf_with_readable_text(self):
        from pypdf import PdfReader

        data = build(_sample(), "pdf").getvalue()
        self.assertTrue(data.startswith(b"%PDF-"))
        text = "".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(data)).pages)
        self.assertIn("Damper clearance", text)

    def test_long_values_wrap_rather_than_running_off_the_page(self):
        """Raw strings in a reportlab table do not wrap - they overflow the
        column and print past the page edge, losing content with no error at
        all. Every cell is a Paragraph for exactly that reason."""
        from pypdf import PdfReader

        long_statement = (
            "The smoke barrier assembly indicated on sheet A-101 is not dimensioned "
            "at the duct penetration, and the referenced detail does not state the "
            "required clearance to combustible construction anywhere on the sheet."
        )
        document = ExportDocument(
            title="Wrap", tables=[ExportTable("T", ["Statement"], [[long_statement]])],
        )
        data = build(document, "pdf").getvalue()
        text = "".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(data)).pages)
        # The tail of the sentence must survive, not just its opening.
        self.assertIn("combustible construction", text.replace("\n", " "))

    def test_an_empty_export_still_produces_a_valid_file(self):
        for export_format in SUPPORTED_FORMATS:
            with self.subTest(export_format=export_format):
                data = build(ExportDocument(title="Nothing yet"), export_format).getvalue()
                self.assertGreater(len(data), 0)
                if export_format == "pdf":
                    self.assertTrue(data.startswith(b"%PDF-"))
                else:
                    # docx and xlsx are both zip containers.
                    self.assertTrue(zipfile.is_zipfile(io.BytesIO(data)))

    def test_a_title_excel_would_reject_is_handled(self):
        """Sheet titles are capped at 31 characters and may not contain
        []:*?/\\ - a real Investigation title breaks both rules routinely, and
        openpyxl raises rather than truncating."""
        from openpyxl import load_workbook

        document = ExportDocument(title="T", tables=[ExportTable(
            "Smoke control: dampers / clearances [A-101] - a very long title indeed",
            ["A"], [["1"]],
        )])
        book = load_workbook(build(document, "xlsx"))
        self.assertTrue(book.sheetnames)
        self.assertLessEqual(len(book.sheetnames[0]), 31)


class TheRouteDeliversThemTests(unittest.TestCase):
    def setUp(self):
        import app as app_module
        from models import User, db

        self.tmp = Path(tempfile.mkdtemp(prefix="beehive_test_export_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp)
        self.project_id = "proj-export"

        with self.flask_app.app_context():
            db.session.add(User(
                username="ex", password_hash=generate_password_hash("x"), role="admin",
            ))
            db.session.commit()

        RequirementsRegistry(self.tmp).save(ParsedDocument(
            project_id=self.project_id, filename="founding.docx",
            ingested_at="2026-01-01T00:00:00+00:00",
            requirements=[RequirementItem(
                id="i1", text="The system shall do a thing.",
                category="other", confidence=0.6, source_line=1,
            )],
        ))
        self.client = self.flask_app.test_client()
        with self.client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "ex"
            sess["role"] = "admin"

        self.store = CaseWorkspaceStore(self.tmp)
        self.workspace = self.store.get_or_create(self.project_id)
        self.store.create_case(self.workspace, title="Sill support", objective="…", created_by="ex")

    def test_every_kind_and_format_downloads(self):
        for kind in ("project", "findings", "requirements", "investigations"):
            for export_format in SUPPORTED_FORMATS:
                with self.subTest(kind=kind, export_format=export_format):
                    response = self.client.get(
                        f"/projects/{self.project_id}/workspace/export/{kind}.{export_format}"
                    )
                    self.assertEqual(response.status_code, 200)
                    self.assertIn("attachment", response.headers.get("Content-Disposition", ""))
                    self.assertGreater(len(response.data), 0)

    def test_the_exported_content_is_really_this_project(self):
        from openpyxl import load_workbook

        response = self.client.get(
            f"/projects/{self.project_id}/workspace/export/investigations.xlsx"
        )
        book = load_workbook(io.BytesIO(response.data))
        values = [
            str(cell.value)
            for sheet in book.worksheets
            for row in sheet.iter_rows()
            for cell in row
            if cell.value is not None
        ]
        self.assertIn("Sill support", values)

    def test_an_unknown_kind_or_format_is_refused(self):
        for path in ("secrets.xlsx", "findings.exe", "findings.docx.exe"):
            with self.subTest(path=path):
                response = self.client.get(
                    f"/projects/{self.project_id}/workspace/export/{path}"
                )
                self.assertEqual(response.status_code, 404)

    def test_it_requires_a_signed_in_reviewer(self):
        anon = self.flask_app.test_client()
        response = anon.get(f"/projects/{self.project_id}/workspace/export/findings.pdf")
        self.assertIn(response.status_code, (302, 401, 403))


class BoundariesTests(unittest.TestCase):
    def test_it_goes_through_the_existing_export_gate(self):
        """ACTION_EXPORT, via the same helper the RFI export uses. A second,
        weaker export door beside the governed one would undo the gate."""
        route = (ROOT / "routes" / "workspace.py").read_text(encoding="utf-8")
        block = route[route.index("def export_document("):]
        block = block[: block.index("return send_file")]
        self.assertIn("_require_export_allowed", block)

    def test_the_writers_cannot_reach_the_store_at_all(self):
        """An export is a VIEW of governed state. The module that writes files
        never receives the store, so it cannot read beyond what it was handed
        or write anything back."""
        # Scanned against CODE, not prose: this module's docstring says "no
        # builder receives the store", which a bare substring check happily
        # accepts as evidence of the opposite. And `save(` is too loose -
        # workbook.save(buffer) writes to an in-memory BytesIO, which is the
        # whole design. The real property is that it can reach neither the
        # store nor the filesystem.
        code = _code_of(ROOT / "services" / "document_export.py")
        for forbidden in ("CaseWorkspaceStore", "call_llm_json", "open(",
                          "write_text", "write_bytes", "REGISTRY_STORE_PATH"):
            with self.subTest(token=forbidden):
                self.assertNotIn(forbidden, code)

    def test_no_builder_takes_the_store_as_an_argument(self):
        """Structural: it cannot read beyond what it was handed."""
        import inspect

        from services import document_export

        for name in ("build_docx", "build_xlsx", "build_pdf", "build"):
            with self.subTest(builder=name):
                parameters = inspect.signature(getattr(document_export, name)).parameters
                self.assertNotIn("store", parameters)
                self.assertNotIn("workspace", parameters)

    def test_nothing_is_summarised_or_reworded_on_the_way_out(self):
        """A document that quietly editorialised would be evidence laundering -
        a reader has no way to tell the difference."""
        code = _code_of(ROOT / "services" / "document_export.py").lower()
        for forbidden in ("anthropic", "llm", "summar", "rewrite"):
            with self.subTest(token=forbidden):
                self.assertNotIn(forbidden, code)

    def test_the_new_dependency_is_pinned_and_justified(self):
        requirements = (ROOT / "requirements.txt").read_text(encoding="utf-8")
        self.assertIn("reportlab==", requirements)
        self.assertIn("dependency_fit.py", requirements)

    def test_the_capability_describes_itself_from_the_start(self):
        """Three reports this session were "the feature exists and I cannot
        reach it". This one is registered in the same change that builds it."""
        from services.capability_registry import find_capability_by_phrase

        for phrase in ("export", "download", "pdf", "excel"):
            with self.subTest(phrase=phrase):
                capability = find_capability_by_phrase(f"how do i {phrase} this")
                self.assertIsNotNone(capability)
                self.assertEqual(capability.key, "export_document")


if __name__ == "__main__":
    unittest.main()
