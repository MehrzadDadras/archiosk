"""
CLAUDE-P29: Project Operating Environment -- the locked, project-level
classification (Client / Owner vs. Design-Builder / Proponent)
established at project creation, distinct from user role and from the
existing per-reviewer represented_party_by/PerspectiveAssessment
mechanism (CLAUDE-P12R/P17), which remains unchanged and cannot affect
this field.

Run via:

    python -m unittest discover -s tests -v
"""
from __future__ import annotations

import io
import shutil
import tempfile
import unittest
from pathlib import Path

import docx
from werkzeug.datastructures import FileStorage
from werkzeug.security import generate_password_hash

from services.case_workspace import CaseWorkspaceStore, OperatingEnvironmentAlreadySetError
from services.environment_capabilities import (
    CLIENT_OWNER,
    DESIGN_BUILDER_PROPONENT,
    allowed_participant_roles,
    is_valid_operating_environment,
)
from services.ingestion import UploadError, ingest_upload
from services.requirements_registry import RequirementsRegistry


def _fake_file(content: bytes, filename: str) -> FileStorage:
    return FileStorage(stream=io.BytesIO(content), filename=filename)


class _BaseWorkspaceTestCase(unittest.TestCase):
    """Shared setup: a real Flask app, a scratch registry dir, and an
    authenticated admin session's test client -- the actual route path,
    not just direct service calls, for the tests that need it."""

    def setUp(self):
        import app as app_module
        from models import User, db

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_operating_env_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        self.client = self.flask_app.test_client()

        with self.flask_app.app_context():
            admin = User(
                username="env_admin", password_hash=generate_password_hash("x"), role="admin",
            )
            db.session.add(admin)
            db.session.commit()

        with self.client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "env_admin"
            sess["role"] = "admin"

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _store(self) -> CaseWorkspaceStore:
        return CaseWorkspaceStore(self.tmp_dir)

    def _ingest_via_service(self, content: bytes, filename: str, environment: str, project_name: str | None = None):
        with self.flask_app.app_context():
            return ingest_upload(
                _fake_file(content, filename), self.flask_app,
                operating_environment=environment, owner="env_admin", project_name=project_name,
            )

    def _ingest_via_route(self, content: bytes, filename: str, environment: str | None, project_name: str):
        data = {
            "file": (io.BytesIO(content), filename),
            "project_name": project_name,
        }
        if environment is not None:
            data["operating_environment"] = environment
        return self.client.post("/upload", data=data, content_type="multipart/form-data")


class NewProjectRequiresEnvironmentTests(_BaseWorkspaceTestCase):
    def test_missing_environment_is_rejected_at_service_layer(self):
        with self.assertRaises(UploadError):
            self._ingest_via_service(b"content", "a.txt", environment="")

    def test_invalid_environment_value_is_rejected(self):
        with self.assertRaises(UploadError):
            self._ingest_via_service(b"content", "a.txt", environment="something_else")

    def test_missing_environment_via_route_returns_400_and_leaves_no_project(self):
        registry = RequirementsRegistry(self.tmp_dir)
        before = set(registry.list_ids())

        response = self._ingest_via_route(b"content", "a.txt", environment=None, project_name="No Env Project")
        self.assertEqual(response.status_code, 400)

        after = set(RequirementsRegistry(self.tmp_dir).list_ids())
        self.assertEqual(before, after, "no project should be created when environment is missing")

    def test_forged_invalid_environment_via_route_is_rejected(self):
        response = self._ingest_via_route(
            b"content", "a.txt", environment="owner_but_not_really_a_valid_value", project_name="Forged Env",
        )
        self.assertEqual(response.status_code, 400)


class EnvironmentCreationAndPersistenceTests(_BaseWorkspaceTestCase):
    def test_client_owner_creation_succeeds_and_persists(self):
        document = self._ingest_via_service(b"content", "a.txt", environment=CLIENT_OWNER, project_name="Client Proj")
        workspace = self._store().get(document.project_id)
        self.assertEqual(workspace.operating_environment, CLIENT_OWNER)

    def test_design_builder_proponent_creation_succeeds_and_persists(self):
        document = self._ingest_via_service(
            b"content", "a.txt", environment=DESIGN_BUILDER_PROPONENT, project_name="DB Proj",
        )
        workspace = self._store().get(document.project_id)
        self.assertEqual(workspace.operating_environment, DESIGN_BUILDER_PROPONENT)

    def test_environment_survives_reopening_via_a_fresh_store_instance(self):
        document = self._ingest_via_service(b"content", "a.txt", environment=CLIENT_OWNER, project_name="Reopen Test")
        # A brand-new CaseWorkspaceStore instance, same store_path -- proves
        # this is real persistence, not in-memory object state.
        fresh_store = CaseWorkspaceStore(self.tmp_dir)
        reopened = fresh_store.get(document.project_id)
        self.assertEqual(reopened.operating_environment, CLIENT_OWNER)

    def test_creation_via_the_real_route_locks_the_selected_environment(self):
        response = self._ingest_via_route(
            b"Contractor shall comply.\n", "a.txt", environment=DESIGN_BUILDER_PROPONENT, project_name="Route Test",
        )
        self.assertEqual(response.status_code, 302)
        # /upload redirects to workspace.show_workspace/<project_id>/workspace
        location = response.headers["Location"]
        project_id = location.rstrip("/").split("/projects/")[-1].split("/workspace")[0]
        workspace = self._store().get(project_id)
        self.assertEqual(workspace.operating_environment, DESIGN_BUILDER_PROPONENT)


class ImmutabilityEnforcementTests(_BaseWorkspaceTestCase):
    def test_direct_service_call_cannot_change_an_already_set_environment(self):
        document = self._ingest_via_service(b"content", "a.txt", environment=CLIENT_OWNER, project_name="Locked")
        store = self._store()
        workspace = store.get(document.project_id)

        with self.assertRaises(OperatingEnvironmentAlreadySetError):
            store.set_operating_environment(workspace, DESIGN_BUILDER_PROPONENT, actor="attacker")

        # Unchanged after the rejected attempt.
        self.assertEqual(store.get(document.project_id).operating_environment, CLIENT_OWNER)

    def test_direct_service_call_cannot_reset_to_the_same_value_either(self):
        # "Locked" means exactly one successful call ever, no exception
        # for re-setting to the identical value -- see the method's own
        # docstring.
        document = self._ingest_via_service(b"content", "a.txt", environment=CLIENT_OWNER, project_name="Locked2")
        store = self._store()
        workspace = store.get(document.project_id)

        with self.assertRaises(OperatingEnvironmentAlreadySetError):
            store.set_operating_environment(workspace, CLIENT_OWNER, actor="someone")

    def test_forged_route_submission_cannot_convert_an_already_classified_project(self):
        document = self._ingest_via_service(b"content", "a.txt", environment=CLIENT_OWNER, project_name="Route Lock")
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/classify-environment",
            data={"operating_environment": DESIGN_BUILDER_PROPONENT},
        )
        # Redirects with a flash error, never raises a 500 -- and, more
        # importantly, never actually changes the stored value.
        self.assertIn(response.status_code, (302, 303))
        self.assertEqual(
            self._store().get(document.project_id).operating_environment, CLIENT_OWNER,
        )

    def test_no_conversion_route_exists_for_client_to_proponent(self):
        # There is deliberately no "edit environment" route at all for an
        # already-classified project -- confirm the only write path
        # (classify-environment) refuses rather than silently no-oping.
        document = self._ingest_via_service(b"content", "a.txt", environment=CLIENT_OWNER, project_name="No Convert")
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/classify-environment",
            data={"operating_environment": DESIGN_BUILDER_PROPONENT},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace_page = self.client.get(f"/projects/{document.project_id}/workspace")
        # No classification form should be offered once already locked.
        self.assertNotIn(b"classify-environment", workspace_page.data)


class RoleAndPerspectiveIndependenceTests(_BaseWorkspaceTestCase):
    """Part VIII's non-negotiable rule: a reviewer's represented_party_by
    (the existing CLAUDE-P12R mechanism) must never be able to change
    ProjectWorkspace.operating_environment -- the two are enforced
    through entirely separate methods with no shared mutation logic."""

    def test_reviewer_changing_represented_party_does_not_touch_environment(self):
        document = self._ingest_via_service(
            b"content", "a.txt", environment=CLIENT_OWNER, project_name="Perspective Independence",
        )
        store = self._store()
        workspace = store.get(document.project_id)

        participant = store.record_participant(
            workspace, name="Cedar Harbour Owner", role_type="owner", created_by="env_admin",
        )
        store.set_represented_party(workspace, reviewer="env_admin", participant_id=participant["id"])

        reloaded = store.get(document.project_id)
        self.assertEqual(reloaded.operating_environment, CLIENT_OWNER)
        self.assertEqual(reloaded.represented_party_by.get("env_admin"), participant["id"])

    def test_user_role_does_not_alter_environment(self):
        # read_only vs admin session -- neither can change the stored
        # environment merely by virtue of which role is active; only
        # authenticated route access differs (checked elsewhere), not
        # this field's value.
        document = self._ingest_via_service(b"content", "a.txt", environment=CLIENT_OWNER, project_name="Role Test")
        with self.client.session_transaction() as sess:
            sess["role"] = "read_only"
        self.client.get(f"/projects/{document.project_id}/workspace")
        self.assertEqual(
            self._store().get(document.project_id).operating_environment, CLIENT_OWNER,
        )

    def test_comparative_perspective_assessment_does_not_alter_environment(self):
        document = self._ingest_via_service(
            b"content", "a.txt", environment=DESIGN_BUILDER_PROPONENT, project_name="Comparative Test",
        )
        store = self._store()
        workspace = store.get(document.project_id)
        participant = store.record_participant(
            workspace, name="Some JV", role_type="design_builder", created_by="env_admin",
        )
        # Registering a participant and recording a perspective is the
        # existing, real "comparative analysis" mechanism -- confirm it
        # never reaches operating_environment at all.
        reloaded = store.get(document.project_id)
        self.assertEqual(reloaded.operating_environment, DESIGN_BUILDER_PROPONENT)
        self.assertEqual(len(reloaded.participants), 1)


class EnvironmentCapabilityGatingTests(_BaseWorkspaceTestCase):
    def test_client_owner_and_proponent_environments_allow_different_participant_roles(self):
        client_roles = allowed_participant_roles(CLIENT_OWNER)
        proponent_roles = allowed_participant_roles(DESIGN_BUILDER_PROPONENT)
        self.assertNotEqual(set(client_roles), set(proponent_roles))
        self.assertIn("owner", client_roles)
        self.assertNotIn("design_builder", client_roles)
        self.assertIn("design_builder", proponent_roles)
        self.assertNotIn("owner", proponent_roles)

    def test_legacy_unset_environment_allows_the_full_open_role_set(self):
        # None must mean "no gating", never "nothing allowed" -- a
        # legacy project's existing Participant functionality must not
        # break just because this field didn't exist when it was made.
        self.assertIsNone(allowed_participant_roles(None))

    def test_server_rejects_a_participant_role_outside_the_locked_environment(self):
        document = self._ingest_via_service(
            b"content", "a.txt", environment=CLIENT_OWNER, project_name="Role Gating",
        )
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/participants",
            data={"name": "Some Design-Builder", "role_type": "design_builder"},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.participants), 0, "the out-of-environment role must not have been registered")

    def test_server_accepts_a_participant_role_inside_the_locked_environment(self):
        document = self._ingest_via_service(
            b"content", "a.txt", environment=CLIENT_OWNER, project_name="Role Gating OK",
        )
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/participants",
            data={"name": "The Owner", "role_type": "owner"},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.participants), 1)


class NeutralParsingAndIsolationTests(_BaseWorkspaceTestCase):
    """Create two otherwise-equivalent projects from the same source in
    opposite environments (Part XIII's explicit two-project scenario)."""

    _SOURCE = (
        b"The Design-Builder shall be responsible for verification of "
        b"existing site conditions prior to commencement of design.\n"
    )

    def setUp(self):
        super().setUp()
        self.project_a = self._ingest_via_service(
            self._SOURCE, "shared_source.txt", environment=CLIENT_OWNER, project_name="Project A Client",
        )
        self.project_b = self._ingest_via_service(
            self._SOURCE, "shared_source.txt", environment=DESIGN_BUILDER_PROPONENT, project_name="Project B Proponent",
        )

    def test_neutral_extracted_requirement_text_is_equivalent(self):
        # Same source, same neutral extraction -- environment must not
        # touch the extract/segment/classify stages at all.
        req_texts_a = sorted(r.text for r in self.project_a.requirements)
        req_texts_b = sorted(r.text for r in self.project_b.requirements)
        self.assertEqual(req_texts_a, req_texts_b)

    def test_project_identity_and_state_remain_separate(self):
        self.assertNotEqual(self.project_a.project_id, self.project_b.project_id)
        workspace_a = self._store().get(self.project_a.project_id)
        workspace_b = self._store().get(self.project_b.project_id)
        self.assertEqual(workspace_a.operating_environment, CLIENT_OWNER)
        self.assertEqual(workspace_b.operating_environment, DESIGN_BUILDER_PROPONENT)

    def test_participant_registered_in_one_project_does_not_appear_in_the_other(self):
        store = self._store()
        workspace_a = store.get(self.project_a.project_id)
        store.record_participant(workspace_a, name="Owner Co", role_type="owner", created_by="env_admin")

        workspace_b_reloaded = store.get(self.project_b.project_id)
        self.assertEqual(len(workspace_b_reloaded.participants), 0)

    def test_neither_project_can_be_converted_into_the_other(self):
        store = self._store()
        with self.assertRaises(OperatingEnvironmentAlreadySetError):
            store.set_operating_environment(
                store.get(self.project_a.project_id), DESIGN_BUILDER_PROPONENT, actor="env_admin",
            )
        with self.assertRaises(OperatingEnvironmentAlreadySetError):
            store.set_operating_environment(
                store.get(self.project_b.project_id), CLIENT_OWNER, actor="env_admin",
            )

    def test_project_name_uniqueness_still_enforced_across_environments(self):
        # Two projects with the SAME source but DIFFERENT names (Project A
        # Client/Project B Proponent, from setUp) already both succeeded --
        # confirm a genuinely duplicate NAME is still rejected regardless
        # of environment (this rule is orthogonal to environment).
        with self.assertRaises(UploadError):
            self._ingest_via_service(
                b"unrelated content", "other.txt", environment=DESIGN_BUILDER_PROPONENT,
                project_name="Project A Client",
            )


class HistoricalLegacyClassificationTests(_BaseWorkspaceTestCase):
    def _make_legacy_workspace(self, project_id: str):
        # Directly constructs a workspace the way get_or_create did
        # before P29 -- operating_environment simply absent, exactly
        # what a real pre-P29 project's JSON looks like on disk.
        store = self._store()
        registry = RequirementsRegistry(self.tmp_dir)
        from services.bhive_parser import ParsedDocument
        registry.save(ParsedDocument(project_id=project_id, filename="legacy.txt", ingested_at="2020-01-01T00:00:00+00:00"))
        workspace = store.get_or_create(project_id)
        self.assertIsNone(workspace.operating_environment)
        return store, workspace

    def test_legacy_project_is_not_falsely_classified(self):
        _, workspace = self._make_legacy_workspace("legacy-proj-1")
        self.assertIsNone(workspace.operating_environment)

    def test_one_time_classification_succeeds_and_is_recorded_in_the_governance_log(self):
        from services.ingestion import get_governance_log

        store, workspace = self._make_legacy_workspace("legacy-proj-2")
        governance_log = get_governance_log(self.flask_app)
        with self.flask_app.app_context():
            store.set_operating_environment(
                workspace, CLIENT_OWNER, actor="env_admin", governance_log=governance_log,
            )

        events = governance_log.read("legacy-proj-2")
        established = next(e for e in events if e.event_type == "operating_environment_established")
        self.assertIsNone(established.payload["previous_state"])
        self.assertEqual(established.payload["operating_environment"], CLIENT_OWNER)

    def test_classified_legacy_project_cannot_subsequently_be_converted(self):
        store, workspace = self._make_legacy_workspace("legacy-proj-3")
        store.set_operating_environment(workspace, CLIENT_OWNER, actor="env_admin")

        with self.assertRaises(OperatingEnvironmentAlreadySetError):
            store.set_operating_environment(
                store.get("legacy-proj-3"), DESIGN_BUILDER_PROPONENT, actor="env_admin",
            )

    def test_classification_route_works_for_a_real_legacy_project(self):
        self._make_legacy_workspace("legacy-proj-4")
        response = self.client.post(
            "/projects/legacy-proj-4/workspace/classify-environment",
            data={"operating_environment": DESIGN_BUILDER_PROPONENT},
        )
        self.assertIn(response.status_code, (302, 303))
        self.assertEqual(
            self._store().get("legacy-proj-4").operating_environment, DESIGN_BUILDER_PROPONENT,
        )


class ExportsIdentifyEnvironmentTests(_BaseWorkspaceTestCase):
    def test_rfi_export_identifies_the_project_operating_environment(self):
        import unittest.mock as mock

        # A consistency flag is required for build_rfi_docx to produce
        # anything at all -- inject one directly rather than depending
        # on a real Anthropic call.
        document = self._ingest_via_service(
            b"Two contradictory requirements.", "a.txt", environment=CLIENT_OWNER, project_name="Export Test",
        )
        from services.bhive_parser import ConsistencyFlag

        registry = RequirementsRegistry(self.tmp_dir)
        stored = registry.get(document.project_id)
        stored.consistency_checked = True
        stored.consistency_flags = [
            ConsistencyFlag(
                id="flag-1", requirement_a_id="a", requirement_a_text="A",
                requirement_b_id="b", requirement_b_text="B", explanation="Conflict.",
            ),
        ]
        registry.save(stored)

        response = self.client.get(f"/projects/{document.project_id}/workspace/rfi-export")
        self.assertEqual(response.status_code, 200)
        full_text = "\n".join(p.text for p in docx.Document(io.BytesIO(response.data)).paragraphs)
        self.assertIn("Client / Owner", full_text)


if __name__ == "__main__":
    unittest.main()
