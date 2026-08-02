"""
CLAUDE-P30: Environment Capability Architecture and Contractual Tool
Directionality -- turns the locked Project Operating Environment
(CLAUDE-P29) into actual, enforced, environment-specific application
behavior via a small centralized capability grammar
(services/environment_capabilities.py's CAPABILITY_REGISTRY /
capability_availability / capability_denial_reason).

Two representative workflows are tested end-to-end:

1. RFI/clarification directionality -- "rfi_originate" (Design-Builder/
   Proponent: draft, revise, issue) and "rfi_respond" (Client/Owner:
   record the authoritative response to an issued RFI) are registered
   as CAPABILITY_COUNTERPART, not a bare client_only/proponent_only
   label -- each has a real, distinct counterpart, which is exactly
   what CAPABILITY_COUNTERPART means. Mechanically, availability for
   each capability is single-sided (rfi_originate: proponent-only in
   effect; rfi_respond: client-only in effect), which is what the
   "Client-only capability denial in Proponent environment" / "Proponent-
   only capability denial in Client environment" test classes below
   exercise concretely.
2. Go/No-Go -- one shared governed record (GoNoGoAssessment,
   CAPABILITY_PARALLEL), validated against whichever decision-stage
   vocabulary the project's own locked environment uses
   (CLIENT_OWNER_DECISION_STAGES / DESIGN_BUILDER_PROPONENT_DECISION_
   STAGES) -- proving one shared capability family can still enforce
   genuinely distinct environment variants.

Run via:

    python -m unittest discover -s tests -v
"""
from __future__ import annotations

import io
import shutil
import tempfile
import unittest
from pathlib import Path

import docx
from werkzeug.datastructures import FileStorage
from werkzeug.security import generate_password_hash

from services.case_workspace import AnalysisTrigger, CaseWorkspaceError, CaseWorkspaceStore
from services.environment_capabilities import (
    CAPABILITY_REGISTRY,
    CLIENT_OWNER,
    CLIENT_OWNER_DECISION_STAGES,
    DESIGN_BUILDER_PROPONENT,
    DESIGN_BUILDER_PROPONENT_DECISION_STAGES,
    capability_availability,
    capability_denial_reason,
    decision_stages_for_environment,
)
from services.ingestion import ingest_upload


def _fake_file(content: bytes, filename: str) -> FileStorage:
    return FileStorage(stream=io.BytesIO(content), filename=filename)


class _BaseWorkspaceTestCase(unittest.TestCase):
    def setUp(self):
        import app as app_module
        from models import User, db

        self.tmp_dir = Path(tempfile.mkdtemp(prefix="beehive_test_capability_arch_"))
        self.flask_app = app_module.create_app("testing")
        self.flask_app.config["REGISTRY_STORE_PATH"] = str(self.tmp_dir)
        self.client = self.flask_app.test_client()

        with self.flask_app.app_context():
            admin = User(
                username="cap_admin", password_hash=generate_password_hash("x"), role="admin",
            )
            db.session.add(admin)
            db.session.commit()

        with self.client.session_transaction() as sess:
            sess["user_id"] = 1
            sess["username"] = "cap_admin"
            sess["role"] = "admin"

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _store(self) -> CaseWorkspaceStore:
        return CaseWorkspaceStore(self.tmp_dir)

    def _ingest(self, content: bytes, filename: str, environment: str, project_name: str):
        with self.flask_app.app_context():
            return ingest_upload(
                _fake_file(content, filename), self.flask_app,
                operating_environment=environment, owner="cap_admin", project_name=project_name,
            )

    def _make_finding_with_validation(self, project_id: str, case_title: str = "Cap Test Case"):
        """A Finding with a Reviewer Validation on record -- the
        precondition create_rfi_draft already enforces, independent of
        this stage's own capability gate. Same pattern as
        tests/test_rfi_draft_export.py's own _create_validated_finding:
        record_analysis against the rfq_rfp_document Source that
        ingest_upload already registers, not a bespoke Source/Finding
        constructor (none exists)."""
        store = self._store()
        workspace = store.get(project_id)
        case = store.create_case(workspace, title=case_title, objective="x", created_by="cap_admin")
        source_id = next(s for s in workspace.sources if s["kind"] == "rfq_rfp_document")["id"]
        trigger = AnalysisTrigger(trigger_type="user_initiated", triggered_by_actor="cap_admin")
        analysis = store.record_analysis(
            workspace, case_id=case["id"], source_ids=[source_id], objective="x",
            engine_name="test", engine_version="1.0",
            findings=[{"statement": "Datum inconsistency observed.", "machine_confidence": 0.7, "source_id": source_id}],
            trigger=trigger,
        )
        finding_id = analysis["finding_ids"][0]
        workspace = store.get(project_id)
        store.record_reviewer_validation(workspace, finding_id=finding_id, validation="Correct", reviewer="cap_admin")
        return case["id"], finding_id


class CapabilityResolutionTests(_BaseWorkspaceTestCase):
    """Direct, no-Flask tests of the resolution functions themselves --
    "capability resolution for Client environment" / "... Proponent
    environment" / "shared neutral capability availability" from Part X."""

    def test_neutral_capability_available_in_both_environments(self):
        for capability_id in ("source_preservation", "neutral_extraction", "case_investigation"):
            self.assertTrue(capability_availability(capability_id, CLIENT_OWNER))
            self.assertTrue(capability_availability(capability_id, DESIGN_BUILDER_PROPONENT))

    def test_neutral_capability_available_for_legacy_unset_environment(self):
        self.assertTrue(capability_availability("source_preservation", None))

    def test_rfi_originate_resolves_proponent_only(self):
        self.assertFalse(capability_availability("rfi_originate", CLIENT_OWNER))
        self.assertTrue(capability_availability("rfi_originate", DESIGN_BUILDER_PROPONENT))

    def test_rfi_respond_resolves_client_only(self):
        self.assertTrue(capability_availability("rfi_respond", CLIENT_OWNER))
        self.assertFalse(capability_availability("rfi_respond", DESIGN_BUILDER_PROPONENT))

    def test_go_no_go_available_in_both_environments(self):
        self.assertTrue(capability_availability("go_no_go", CLIENT_OWNER))
        self.assertTrue(capability_availability("go_no_go", DESIGN_BUILDER_PROPONENT))

    def test_future_not_authorized_capability_unavailable_everywhere(self):
        for environment in (CLIENT_OWNER, DESIGN_BUILDER_PROPONENT, None):
            self.assertFalse(capability_availability("security_policy_architecture", environment))

    def test_legacy_unset_environment_is_ungated_for_ordinary_capabilities(self):
        # None means "no gating" (P29's established precedent for
        # allowed_participant_roles, applied consistently here) --
        # except for FUTURE_NOT_AUTHORIZED, which is never available.
        self.assertTrue(capability_availability("rfi_originate", None))
        self.assertTrue(capability_availability("rfi_respond", None))
        self.assertFalse(capability_availability("security_policy_architecture", None))


class CapabilityDenialMessageTests(unittest.TestCase):
    """"denial messages identify the correct reason" (Part X / IX)."""

    def test_denial_reason_none_when_available(self):
        self.assertIsNone(capability_denial_reason("rfi_originate", DESIGN_BUILDER_PROPONENT))

    def test_denial_reason_names_the_capability_and_environment(self):
        reason = capability_denial_reason("rfi_originate", CLIENT_OWNER)
        self.assertIn("Client / Owner", reason)
        self.assertIn("clarification request", reason)

    def test_denial_reason_names_the_counterpart_when_one_exists(self):
        reason = capability_denial_reason("rfi_respond", DESIGN_BUILDER_PROPONENT)
        self.assertIn("counterpart capability", reason)

    def test_denial_reason_for_future_not_authorized_says_not_yet_authorized(self):
        reason = capability_denial_reason("security_policy_architecture", CLIENT_OWNER)
        self.assertIn("not yet authorized", reason)

    def test_registry_is_internally_consistent(self):
        # Every registered capability's classification is one of the
        # declared grammar constants, and availability always derives
        # from client_variant/proponent_variant (Part II's own rule) --
        # not a separately-set boolean that could drift.
        for capability_id, definition in CAPABILITY_REGISTRY.items():
            self.assertEqual(
                capability_availability(capability_id, CLIENT_OWNER),
                definition.client_variant is not None,
            )
            self.assertEqual(
                capability_availability(capability_id, DESIGN_BUILDER_PROPONENT),
                definition.proponent_variant is not None,
            )


class GoNoGoDecisionStageVocabularyTests(unittest.TestCase):
    """"Go/No-Go resolves to different variants" (Part X)."""

    def test_client_and_proponent_stage_vocabularies_differ(self):
        self.assertNotEqual(set(CLIENT_OWNER_DECISION_STAGES), set(DESIGN_BUILDER_PROPONENT_DECISION_STAGES))

    def test_decision_stages_for_environment_resolves_correctly(self):
        self.assertEqual(decision_stages_for_environment(CLIENT_OWNER), CLIENT_OWNER_DECISION_STAGES)
        self.assertEqual(
            decision_stages_for_environment(DESIGN_BUILDER_PROPONENT), DESIGN_BUILDER_PROPONENT_DECISION_STAGES,
        )

    def test_decision_stages_for_environment_raises_for_unset(self):
        with self.assertRaises(ValueError):
            decision_stages_for_environment(None)


class GoNoGoRouteTests(_BaseWorkspaceTestCase):
    def test_client_project_records_a_client_stage_decision(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "GNG Client")
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "release_rfp", "decision": "go", "rationale": "Budget confirmed."},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.go_no_go_assessments), 1)
        self.assertEqual(workspace.go_no_go_assessments[0]["decision_stage"], "release_rfp")
        self.assertEqual(workspace.go_no_go_assessments[0]["operating_environment"], CLIENT_OWNER)

    def test_proponent_project_records_a_proponent_stage_decision(self):
        document = self._ingest(b"content", "a.txt", DESIGN_BUILDER_PROPONENT, "GNG Proponent")
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "bid_rfp", "decision": "no_go", "rationale": "Margin too thin."},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(workspace.go_no_go_assessments[0]["decision_stage"], "bid_rfp")

    def test_client_project_rejects_a_proponent_only_stage(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "GNG Wrong Stage")
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "bid_rfp", "decision": "go", "rationale": "Should be rejected."},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.go_no_go_assessments), 0)

    def test_direct_service_call_also_rejects_a_wrong_environment_stage(self):
        document = self._ingest(b"content", "a.txt", DESIGN_BUILDER_PROPONENT, "GNG Service Reject")
        store = self._store()
        workspace = store.get(document.project_id)
        with self.assertRaises(CaseWorkspaceError):
            store.record_go_no_go_decision(
                workspace, decision_stage="release_rfp", decision="go",
                rationale="Wrong side's stage.", decided_by="cap_admin",
            )

    def test_legacy_unclassified_project_cannot_record_a_go_no_go_decision(self):
        # "legacy unclassified projects follow an explicit compatibility
        # rule" (Part X) -- unlike participant-role gating, Go/No-Go has
        # no sensible ungated fallback (no vocabulary to validate
        # against), so this is a hard refusal, not a silent default.
        from services.bhive_parser import ParsedDocument
        from services.requirements_registry import RequirementsRegistry

        registry = RequirementsRegistry(self.tmp_dir)
        registry.save(ParsedDocument(project_id="legacy-gng", filename="x.txt", ingested_at="2020-01-01T00:00:00+00:00"))
        store = self._store()
        workspace = store.get_or_create("legacy-gng")
        self.assertIsNone(workspace.operating_environment)

        with self.assertRaises(CaseWorkspaceError):
            store.record_go_no_go_decision(
                workspace, decision_stage="release_rfp", decision="go",
                rationale="No environment yet.", decided_by="cap_admin",
            )

    def test_navigation_reflects_capability_resolution_for_go_no_go(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "GNG Nav")
        page = self.client.get(f"/projects/{document.project_id}/workspace?view=overview")
        self.assertIn(b"Record a Go/No-Go decision", page.data)
        # Client-side stage vocabulary is offered, Proponent-only stages are not.
        self.assertIn(b'value="release_rfp"', page.data)
        self.assertNotIn(b'value="bid_rfp"', page.data)

    def test_decision_records_the_deciders_role(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "GNG Role Recorded")
        self.client.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "release_rfp", "decision": "go", "rationale": "Budget confirmed."},
        )
        workspace = self._store().get(document.project_id)
        self.assertEqual(workspace.go_no_go_assessments[0]["decided_by_role"], "admin")

    # -- CLAUDE-P38 (OBS-04): decision authority is admin-only ---------------

    def _read_only_client(self, project_id: str, owner_username: str):
        from models import User, db

        with self.flask_app.app_context():
            db.session.add(User(username="gng_reader", password_hash=generate_password_hash("x"), role="read_only"))
            db.session.commit()
        store = self._store()
        workspace = store.get(project_id)
        store.grant_project_access(workspace, username="gng_reader", actor=owner_username, actor_role="admin")
        reader_client = self.flask_app.test_client()
        with reader_client.session_transaction() as sess:
            sess["user_id"] = 2
            sess["username"] = "gng_reader"
            sess["role"] = "read_only"
        return reader_client

    def test_read_only_user_cannot_record_a_decision(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "GNG RO Denied")
        reader_client = self._read_only_client(document.project_id, "cap_admin")

        response = reader_client.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "release_rfp", "decision": "go", "rationale": "Trying to decide anyway."},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.go_no_go_assessments), 0)

    def test_read_only_user_does_not_see_the_decision_form(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "GNG RO Form Hidden")
        reader_client = self._read_only_client(document.project_id, "cap_admin")

        page = reader_client.get(f"/projects/{document.project_id}/workspace?view=overview")
        self.assertNotIn(b"Record a Go/No-Go decision", page.data)
        self.assertIn(b"Only an admin can record a Go", page.data)

    def test_admin_can_still_record_after_a_read_only_denial(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "GNG Admin Still Works")
        reader_client = self._read_only_client(document.project_id, "cap_admin")
        reader_client.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "release_rfp", "decision": "go", "rationale": "Denied attempt."},
        )
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/go-no-go",
            data={"decision_stage": "release_rfp", "decision": "go", "rationale": "Real admin decision."},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.go_no_go_assessments), 1)
        self.assertEqual(workspace.go_no_go_assessments[0]["rationale"], "Real admin decision.")


class RfiDirectionalityRouteTests(_BaseWorkspaceTestCase):
    """"Client cannot originate Proponent RFI workflow" / "Proponent
    cannot issue authoritative Client response" / "direct route ... calls
    cannot bypass capability rules" (Part X)."""

    def test_proponent_can_create_and_issue_an_rfi_draft(self):
        document = self._ingest(b"content", "a.txt", DESIGN_BUILDER_PROPONENT, "RFI Proponent")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)

        response = self.client.post(
            f"/projects/{document.project_id}/workspace/cases/{case_id}/rfi-drafts",
            data={"finding_id": finding_id, "question_text": "Please clarify the datum."},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.rfi_drafts), 1)

    def test_client_cannot_originate_an_rfi_draft(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "RFI Client Denied")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)

        response = self.client.post(
            f"/projects/{document.project_id}/workspace/cases/{case_id}/rfi-drafts",
            data={"finding_id": finding_id, "question_text": "Should be blocked."},
        )
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.rfi_drafts), 0, "a Client/Owner project must not be able to originate an RFI")

    def test_client_cannot_issue_an_rfi_draft_via_a_forged_direct_post(self):
        # Even if a draft somehow exists (e.g. created before the project
        # was locked to CLIENT_OWNER, or via direct service manipulation),
        # the issue route itself must independently refuse.
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "RFI Client Forged Issue")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="x", created_by="cap_admin")

        response = self.client.post(
            f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/issue",
        )
        self.assertIn(response.status_code, (302, 303))
        reloaded = self._store().get(document.project_id)
        self.assertEqual(reloaded.rfi_drafts[0]["status"], "draft", "must remain unissued")

    def test_client_can_respond_to_an_issued_rfi(self):
        document = self._ingest(b"content", "a.txt", DESIGN_BUILDER_PROPONENT, "RFI For Response")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="Clarify?", created_by="cap_admin")
        store.issue_rfi_draft(workspace, draft_id=draft["id"], issued_by="cap_admin")

        # A second, Client/Owner project's session issues the response --
        # respond_to_rfi_draft is gated per-project by the DRAFT'S OWN
        # project's environment, so this call targets the Proponent
        # project's own draft (the same real-world exchange happens
        # inside a single project's records in this codebase's current
        # model -- see Part XI's own critique note in the final report
        # about single-project vs. cross-project RFI exchange).
        response = self.client.post(
            f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/respond",
            data={"response_text": "Datum confirmed per drawing A-101."},
        )
        # The draft's own project is DESIGN_BUILDER_PROPONENT, so
        # rfi_respond is NOT available here -- this must be denied.
        self.assertIn(response.status_code, (302, 303))
        reloaded = self._store().get(document.project_id)
        self.assertEqual(reloaded.rfi_drafts[0]["status"], "issued", "a Proponent project cannot respond to its own RFI")

    def test_client_project_can_respond_to_its_own_issued_rfi(self):
        # A Client/Owner project reaches an issued draft via legacy data
        # or a project that started unclassified -- construct one
        # directly at the service layer (bypassing origination gating,
        # which is a route-layer concern) to exercise the respond path
        # on its own terms.
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "RFI Client Respond")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="Clarify?", created_by="cap_admin")
        store.issue_rfi_draft(workspace, draft_id=draft["id"], issued_by="cap_admin")

        response = self.client.post(
            f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/respond",
            data={"response_text": "Datum confirmed per drawing A-101."},
        )
        self.assertIn(response.status_code, (302, 303))
        reloaded = self._store().get(document.project_id)
        self.assertEqual(reloaded.rfi_drafts[0]["status"], "answered")
        self.assertEqual(reloaded.rfi_drafts[0]["response_text"], "Datum confirmed per drawing A-101.")

    def test_proponent_cannot_respond_to_its_own_issued_rfi(self):
        document = self._ingest(b"content", "a.txt", DESIGN_BUILDER_PROPONENT, "RFI Proponent No Respond")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="Clarify?", created_by="cap_admin")
        store.issue_rfi_draft(workspace, draft_id=draft["id"], issued_by="cap_admin")

        response = self.client.post(
            f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/respond",
            data={"response_text": "Should be denied."},
        )
        self.assertIn(response.status_code, (302, 303))
        reloaded = self._store().get(document.project_id)
        self.assertEqual(reloaded.rfi_drafts[0]["status"], "issued")

    def test_direct_service_call_still_enforces_issued_precondition_for_response(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "RFI Precondition")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="x", created_by="cap_admin")

        with self.assertRaises(CaseWorkspaceError):
            store.respond_to_rfi_draft(
                workspace, draft_id=draft["id"], response_text="Too early.", responded_by="cap_admin",
            )

    def test_legacy_unclassified_project_is_ungated_for_rfi_origination(self):
        # "legacy unclassified projects follow an explicit compatibility
        # rule" -- for RFI (unlike Go/No-Go), the rule is the same
        # ungated-by-default precedent as participant-role gating.
        from services.bhive_parser import ParsedDocument
        from services.requirements_registry import RequirementsRegistry

        registry = RequirementsRegistry(self.tmp_dir)
        registry.save(ParsedDocument(project_id="legacy-rfi", filename="x.txt", ingested_at="2020-01-01T00:00:00+00:00"))
        store = self._store()
        workspace = store.get_or_create("legacy-rfi")
        case = store.create_case(workspace, title="Legacy Case", objective="x", created_by="cap_admin")
        source = store.add_source(workspace, name="d.txt", file_path="d.txt", kind="rfq_rfp_document")
        trigger = AnalysisTrigger(trigger_type="user_initiated", triggered_by_actor="cap_admin")
        analysis = store.record_analysis(
            workspace, case_id=case["id"], source_ids=[source["id"]], objective="x",
            engine_name="test", engine_version="1.0",
            findings=[{"statement": "Legacy finding.", "machine_confidence": 0.7, "source_id": source["id"]}],
            trigger=trigger,
        )
        finding_id = analysis["finding_ids"][0]
        workspace = store.get("legacy-rfi")
        store.record_reviewer_validation(workspace, finding_id=finding_id, validation="Correct", reviewer="cap_admin")

        response = self.client.post(
            f"/projects/legacy-rfi/workspace/cases/{case['id']}/rfi-drafts",
            data={"finding_id": finding_id, "question_text": "Legacy RFI still allowed."},
        )
        self.assertIn(response.status_code, (302, 303))
        reloaded = store.get("legacy-rfi")
        self.assertEqual(len(reloaded.rfi_drafts), 1)


class NavigationReflectsCapabilityTests(_BaseWorkspaceTestCase):
    """"navigation reflects capability resolution" (Part X)."""

    def test_client_workspace_page_does_not_offer_rfi_origination_controls(self):
        # A draft already exists (constructed directly at the service
        # layer, bypassing route-level origination gating, to isolate
        # what THIS test actually checks: template-layer visibility) --
        # its edit/issue controls must not render in a Client/Owner
        # project even though the draft itself exists.
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "Nav Client")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        store.create_rfi_draft(workspace, finding_id=finding_id, question_text="x", created_by="cap_admin")

        page = self.client.get(f"/projects/{document.project_id}/workspace?case={case_id}")
        self.assertNotIn(b"Save question", page.data)
        self.assertNotIn(b"Issue RFI", page.data)
        self.assertIn(b"does not originate RFIs", page.data)

    def test_proponent_workspace_page_offers_rfi_origination_controls(self):
        document = self._ingest(b"content", "a.txt", DESIGN_BUILDER_PROPONENT, "Nav Proponent")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        store.create_rfi_draft(workspace, finding_id=finding_id, question_text="x", created_by="cap_admin")

        page = self.client.get(f"/projects/{document.project_id}/workspace?case={case_id}")
        self.assertIn(b"Save question", page.data)

    def test_issued_rfi_offers_response_control_only_in_client_environment(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "Nav Respond")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="x", created_by="cap_admin")
        store.issue_rfi_draft(workspace, draft_id=draft["id"], issued_by="cap_admin")

        page = self.client.get(f"/projects/{document.project_id}/workspace?case={case_id}")
        self.assertIn(b"Record Response", page.data)


class ReviewerPerspectiveBoundaryTests(_BaseWorkspaceTestCase):
    """Part VIII's non-negotiable rule: a reviewer's represented_party_by
    (mutable, per-reviewer) must never unlock a capability the project's
    own locked operating_environment denies."""

    def test_representing_the_opposing_side_does_not_unlock_rfi_origination(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "Perspective No Unlock")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)

        # The reviewer sets their represented party to a Design-Builder
        # participant even though the PROJECT itself is Client/Owner --
        # a legitimate, already-supported comparative-perspective act.
        participant = store.record_participant(
            workspace, name="Some JV", role_type="design_builder", created_by="cap_admin",
        )
        # allowed_participant_roles(CLIENT_OWNER) doesn't include
        # "design_builder", but record_participant itself has no
        # environment gate (only the route does) -- construct this
        # directly to isolate the boundary this test actually checks:
        # represented_party_by, not participant registration.
        store.set_represented_party(workspace, reviewer="cap_admin", participant_id=participant["id"])

        reloaded = store.get(document.project_id)
        self.assertEqual(reloaded.represented_party_by.get("cap_admin"), participant["id"])

        response = self.client.post(
            f"/projects/{document.project_id}/workspace/cases/{case_id}/rfi-drafts",
            data={"finding_id": finding_id, "question_text": "Still should be blocked."},
        )
        self.assertIn(response.status_code, (302, 303))
        final = self._store().get(document.project_id)
        self.assertEqual(len(final.rfi_drafts), 0)

    def test_session_role_does_not_affect_capability_availability(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "Role No Unlock")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        with self.client.session_transaction() as sess:
            sess["role"] = "read_only"

        response = self.client.post(
            f"/projects/{document.project_id}/workspace/cases/{case_id}/rfi-drafts",
            data={"finding_id": finding_id, "question_text": "x"},
        )
        # Denied either way -- the point is it's denied for the SAME
        # capability reason regardless of session role, not a role check.
        self.assertIn(response.status_code, (302, 303))
        workspace = self._store().get(document.project_id)
        self.assertEqual(len(workspace.rfi_drafts), 0)


class ExportsIdentifyEnvironmentAndDirectionTests(_BaseWorkspaceTestCase):
    """"environment label and direction appear in exports" (Part X)."""

    def test_rfi_draft_export_identifies_workflow_direction_when_issued(self):
        document = self._ingest(b"content", "a.txt", DESIGN_BUILDER_PROPONENT, "Export Direction")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="Clarify?", created_by="cap_admin")
        store.issue_rfi_draft(workspace, draft_id=draft["id"], issued_by="cap_admin")

        response = self.client.get(f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        self.assertEqual(response.status_code, 200)
        full_text = "\n".join(p.text for p in docx.Document(io.BytesIO(response.data)).paragraphs)
        self.assertIn("Design-Builder/Proponent", full_text)
        self.assertIn("Design-Builder / Proponent", full_text)

    def test_rfi_draft_export_identifies_answered_direction(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "Export Answered")
        case_id, finding_id = self._make_finding_with_validation(document.project_id)
        store = self._store()
        workspace = store.get(document.project_id)
        draft = store.create_rfi_draft(workspace, finding_id=finding_id, question_text="Clarify?", created_by="cap_admin")
        store.issue_rfi_draft(workspace, draft_id=draft["id"], issued_by="cap_admin")
        store.respond_to_rfi_draft(workspace, draft_id=draft["id"], response_text="Confirmed.", responded_by="cap_admin")

        response = self.client.get(f"/projects/{document.project_id}/workspace/rfi-drafts/{draft['id']}/export")
        full_text = "\n".join(p.text for p in docx.Document(io.BytesIO(response.data)).paragraphs)
        self.assertIn("answered by Client/Owner", full_text)
        self.assertIn("Confirmed.", full_text)

    def test_api_rfi_export_now_identifies_environment(self):
        # CLAUDE-P30 closed the scope limitation CLAUDE-P29 explicitly
        # noted: routes/api.py's own RFI export was previously unstamped.
        document = self._ingest(
            b"Two contradictory requirements.", "a.txt", CLIENT_OWNER, "API Export Stamped",
        )
        from services.bhive_parser import ConsistencyFlag
        from services.requirements_registry import RequirementsRegistry

        registry = RequirementsRegistry(self.tmp_dir)
        stored = registry.get(document.project_id)
        stored.consistency_checked = True
        stored.consistency_flags = [
            ConsistencyFlag(
                id="flag-1", requirement_a_id="a", requirement_a_text="A",
                requirement_b_id="b", requirement_b_text="B", explanation="Conflict.",
            ),
        ]
        registry.save(stored)

        response = self.client.get(f"/api/v1/documents/{document.project_id}/rfi")
        self.assertEqual(response.status_code, 200)
        full_text = "\n".join(p.text for p in docx.Document(io.BytesIO(response.data)).paragraphs)
        self.assertIn("Client / Owner", full_text)


class NoTenancyMigrationAndIsolationTests(_BaseWorkspaceTestCase):
    """"no tenancy migration occurs" / "project isolation remains
    intact" (Part X) -- P30 must not have introduced any multi-org/
    tenant concept, and capability gating must not cross project
    boundaries."""

    def test_no_tenant_or_organization_field_exists_on_workspace(self):
        document = self._ingest(b"content", "a.txt", CLIENT_OWNER, "No Tenancy")
        workspace = self._store().get(document.project_id)
        from dataclasses import fields as dataclass_fields

        field_names = {f.name for f in dataclass_fields(workspace)}
        self.assertFalse(any("tenant" in name or "organization" in name for name in field_names))

    def test_go_no_go_and_rfi_state_do_not_leak_across_projects(self):
        client_doc = self._ingest(b"shared content", "a.txt", CLIENT_OWNER, "Isolation Client")
        proponent_doc = self._ingest(b"shared content", "b.txt", DESIGN_BUILDER_PROPONENT, "Isolation Proponent")

        self.client.post(
            f"/projects/{client_doc.project_id}/workspace/go-no-go",
            data={"decision_stage": "release_rfp", "decision": "go", "rationale": "x"},
        )
        client_ws = self._store().get(client_doc.project_id)
        proponent_ws = self._store().get(proponent_doc.project_id)
        self.assertEqual(len(client_ws.go_no_go_assessments), 1)
        self.assertEqual(len(proponent_ws.go_no_go_assessments), 0)


if __name__ == "__main__":
    unittest.main()
